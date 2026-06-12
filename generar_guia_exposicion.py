# -*- coding: utf-8 -*-
"""Genera Guia_Exposicion_SuperAhorro.docx con el guion de la exposicion."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Estilo base
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

ACCENT = RGBColor(0x1B, 0x5E, 0x20)
GRAY = RGBColor(0x66, 0x66, 0x66)


def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_code(code, filename=None):
    if filename:
        p = doc.add_paragraph()
        r = p.add_run("Archivo: " + filename)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, "F2F2F2")
    lines = code.strip("\n").split("\n")
    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line.strip() else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
    doc.add_paragraph()


def add_speaker(name):
    p = doc.add_paragraph()
    r = p.add_run("Lo presenta: " + name)
    r.bold = True
    r.font.color.rgb = ACCENT
    r.font.size = Pt(10.5)


def add_subhead(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)


def add_tag(tipo):
    p = doc.add_paragraph()
    r = p.add_run(tipo)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xFFFFFF >> 16, (0xFFFFFF >> 8) & 0xFF, 0xFFFFFF & 0xFF)
    return p


# ===========================================================================
# PORTADA
# ===========================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("SuperAhorro")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Guia de exposicion - Segunda entrega y etapa final")
r.font.size = Pt(16)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Funcionalidades obligatorias y agregados diferenciales, con explicacion tecnica y codigo real")
r.font.size = Pt(11)
r.font.color.rgb = GRAY

doc.add_paragraph()
integ = doc.add_paragraph()
integ.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = integ.add_run("Integrantes: Bossio / Correa")
r.bold = True

doc.add_page_break()

# ===========================================================================
# INTRODUCCION (ambos)
# ===========================================================================
doc.add_heading("0. Introduccion (ambos integrantes)", level=1)
add_speaker("Bossio y Correa, en conjunto")

doc.add_paragraph(
    "SuperAhorro es una aplicacion Android para registrar y analizar las compras de "
    "supermercado. Permite guardar tickets, ver el historial y las estadisticas de gasto, "
    "comparar precios entre supermercados, recibir promociones y consultar un asistente "
    "con inteligencia artificial sobre el historial de compras."
)
doc.add_paragraph(
    "La app esta hecha en Kotlin con Jetpack Compose para toda la interfaz, sigue el "
    "patron MVVM (Model-View-ViewModel) y usa Firebase como backend: Firebase "
    "Authentication para el login y Cloud Firestore como base de datos. Todas las "
    "operaciones de red y de base de datos se hacen con corrutinas de Kotlin, para no "
    "bloquear la interfaz."
)
doc.add_paragraph(
    "Vamos a mostrar las siete funcionalidades obligatorias de esta entrega "
    "(persistencia de sesion, base de datos, corrutinas, networking, menus y dialogos, "
    "carga real de datos e Intents) y seis agregados diferenciales de la etapa final: "
    "autenticacion biometrica, sincronizacion en la nube, notificaciones, carga "
    "automatica y OCR de tickets con IA, chat sobre el historial y comparador de precios."
)

add_subhead("Aclaracion importante sobre la base de datos local")
doc.add_paragraph(
    "En la primera entrega usabamos Room (SQLite) como base de datos local. Para esta "
    "etapa migramos todo a Cloud Firestore, la base de datos en la nube de Firebase. "
    "Pero el SDK de Firestore para Android mantiene, por defecto, una copia local de los "
    "datos en una base SQLite interna que funciona como cache: toda lectura y escritura "
    "se hace primero contra esa copia local, y despues se sincroniza con el servidor en "
    "segundo plano cuando hay conexion. Por eso, cuando mostremos el repositorio de "
    "compras (CompraRepository), en realidad estamos mostrando a la vez la base de datos "
    "local (la cache offline) y la sincronizacion en la nube (el agregado diferencial)."
)

doc.add_heading("Mapa de funcionalidades y reparto", level=2)

tabla = doc.add_table(rows=1, cols=4)
tabla.style = "Light Grid Accent 1"
hdr = tabla.rows[0].cells
for i, h in enumerate(["#", "Funcionalidad", "Tipo", "Presenta"]):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for run in p.runs:
            run.bold = True

filas = [
    ("1", "Persistencia de sesion (DataStore)", "Obligatorio", "Bossio"),
    ("2", "Autenticacion biometrica", "Extra", "Bossio"),
    ("3", "Base de datos de compras y productos (Firestore)", "Obligatorio", "Bossio"),
    ("4", "Sincronizacion en la nube", "Extra", "Bossio"),
    ("5", "Operaciones con corrutinas", "Obligatorio", "Bossio"),
    ("6", "Menus y dialogos", "Obligatorio", "Bossio"),
    ("7", "Compartir compra por Intent", "Obligatorio", "Correa"),
    ("8", "Networking: API publica de supermercados", "Obligatorio", "Correa"),
    ("9", "Carga real de datos + notificaciones de ofertas", "Obligatorio + Extra", "Correa"),
    ("10", "Carga automatica y OCR de ticket con IA", "Extra", "Correa"),
    ("11", "Chat sobre el historial de compras", "Extra", "Correa"),
    ("12", "Comparativa de precios entre supermercados", "Extra", "Correa"),
]
for fila in filas:
    cells = tabla.add_row().cells
    for i, val in enumerate(fila):
        cells[i].text = val

doc.add_paragraph()
doc.add_paragraph(
    "Quedan 6 temas para cada integrante, con la misma cantidad de explicacion tecnica "
    "y codigo. Esto cubre el pedido del profesor de mostrar al menos un caso con "
    "DataStore (tema 1), uno con base de datos local (tema 3, Firestore + cache offline) "
    "y uno con networking consumiendo una API (tema 8), ademas de varios casos extra "
    "(GroqRepository en los temas 10 y 11 tambien es networking)."
)

doc.add_page_break()

# ===========================================================================
# PARTE 1 - BOSSIO
# ===========================================================================
p = doc.add_paragraph()
r = p.add_run("PARTE 1 - BOSSIO")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = ACCENT
doc.add_paragraph(
    "Sesion, base de datos en la nube, corrutinas, menus y dialogos."
)

# --- Tema 1: DataStore ------------------------------------------------------
doc.add_heading("1. Persistencia local de sesion - DataStore", level=1)
add_tag("OBLIGATORIO")
add_speaker("Bossio")

add_subhead("Que decir")
doc.add_paragraph(
    "Arrancamos con uno de los requisitos obligatorios: la persistencia local de "
    "sesion. Cuando un usuario inicia sesion, no queremos que tenga que volver a "
    "escribir su email y contrasena cada vez que abre la app. Para esto usamos "
    "DataStore, la libreria que reemplaza a SharedPreferences en Android moderno, y "
    "que internamente guarda los datos en un archivo de preferencias usando corrutinas "
    "y Flow."
)
doc.add_paragraph(
    "Lo que guardamos es el UID que nos da Firebase Authentication al loguearse, junto "
    "con el email y el nombre del usuario. Mientras ese UID exista en DataStore, la app "
    "considera que hay una sesion activa y manda directo a la pantalla principal sin "
    "pasar por el login."
)

add_subhead("Como lo hicimos")
doc.add_paragraph(
    "Creamos una clase SessionManager que envuelve un DataStore<Preferences>. "
    "Definimos tres claves de tipo String (user_uid, user_email, user_name) y "
    "expusimos tres Flow que emiten esos valores cada vez que cambian. El metodo "
    "guardarSesion() es una funcion suspend que escribe los tres valores con "
    "dataStore.edit{}, y se llama justo despues de un login exitoso (en "
    "AuthRepository.login). Para cerrar sesion, cerrarSesion() vuelve a escribir un "
    "string vacio en las tres claves."
)

add_code(r'''
private val Context.dataStore: DataStore<Preferences> by preferencesDataStore(name = "session")

class SessionManager(private val context: Context) {

    companion object {
        private val KEY_USER_ID    = stringPreferencesKey("user_uid")
        private val KEY_USER_EMAIL = stringPreferencesKey("user_email")
        private val KEY_USER_NAME  = stringPreferencesKey("user_name")
        const val NO_SESSION = ""
    }

    val userId: Flow<String> = context.dataStore.data.map { prefs ->
        prefs[KEY_USER_ID] ?: NO_SESSION
    }

    suspend fun guardarSesion(userId: String, email: String, nombre: String) {
        context.dataStore.edit { prefs ->
            prefs[KEY_USER_ID]    = userId
            prefs[KEY_USER_EMAIL] = email
            prefs[KEY_USER_NAME]  = nombre
        }
    }

    suspend fun cerrarSesion() {
        context.dataStore.edit { prefs ->
            prefs[KEY_USER_ID]    = NO_SESSION
            prefs[KEY_USER_EMAIL] = ""
            prefs[KEY_USER_NAME]  = ""
        }
    }
}
''', "data/local/SessionManager.kt")

doc.add_paragraph(
    "Caso de uso concreto - \"Guardar usuario logueado\": despues de validar el email "
    "y la contrasena contra Firebase Auth, AuthRepository.login() llama a "
    "session.guardarSesion(uid, usuario.email, usuario.nombre):"
)

add_code(r'''
suspend fun login(email: String, password: String): AuthResult {
    val emailNorm = email.trim().lowercase()
    return try {
        val resultado = auth.signInWithEmailAndPassword(emailNorm, password).await()
        val uid = resultado.user?.uid
            ?: return AuthResult.Error("Email o contraseña incorrectos")

        val usuario = getUsuario(uid) ?: Usuario("", "", emailNorm)
        session.guardarSesion(uid, usuario.email, usuario.nombre)
        AuthResult.Exito(usuario)
    } catch (e: Exception) {
        AuthResult.Error("No se pudo iniciar sesión: ${e.message}")
    }
}
''', "data/repository/AuthRepository.kt")

doc.add_page_break()

# --- Tema 2: Autenticacion biometrica --------------------------------------
doc.add_heading("2. Autenticacion biometrica", level=1)
add_tag("EXTRA")
add_speaker("Bossio")

add_subhead("Que decir")
doc.add_paragraph(
    "Como primer agregado diferencial, implementamos el login por huella digital o "
    "reconocimiento facial. La idea es que, si el usuario ya inicio sesion una vez con "
    "email y contrasena, la proxima vez puede entrar simplemente con su biometria, sin "
    "volver a escribir nada."
)
doc.add_paragraph(
    "Para esto usamos las librerias androidx.biometric. Primero chequeamos si el "
    "dispositivo tiene sensor de huella o face unlock configurado con BiometricManager. "
    "Si es asi, mostramos un boton de 'Ingresar con biometria'. Al tocarlo, se abre el "
    "dialogo nativo del sistema (BiometricPrompt) donde el usuario pone su huella o "
    "cara. Si el sistema valida la biometria, nuestra app confia en esa validacion y "
    "reabre la sesion guardada en DataStore."
)

add_subhead("Como lo hicimos")
doc.add_paragraph(
    "En LoginScreen comprobamos BiometricManager.from(context).canAuthenticate(...) "
    "para decidir si mostrar el boton. Al presionarlo, construimos un BiometricPrompt "
    "con un AuthenticationCallback: cuando el sistema llama a "
    "onAuthenticationSucceeded, invocamos vm.loginConBiometria(). Esa funcion, en "
    "AuthRepository, lee el UID guardado en DataStore (SessionManager), verifica que "
    "Firebase todavia tenga una sesion activa (auth.currentUser), busca los datos del "
    "usuario en Firestore y vuelve a guardar la sesion. Combinamos: biometria del "
    "sistema + sesion local (DataStore) + datos en la nube (Firestore)."
)

add_code(r'''
val puedeUsarBiometria = remember {
    BiometricManager.from(context).canAuthenticate(
        BiometricManager.Authenticators.BIOMETRIC_STRONG or
        BiometricManager.Authenticators.BIOMETRIC_WEAK
    ) == BiometricManager.BIOMETRIC_SUCCESS
}

fun lanzarBiometria() {
    val activity = context as FragmentActivity
    val prompt = BiometricPrompt(
        activity,
        ContextCompat.getMainExecutor(context),
        object : BiometricPrompt.AuthenticationCallback() {
            override fun onAuthenticationSucceeded(result: BiometricPrompt.AuthenticationResult) {
                vm.loginConBiometria(onLoginExitoso)
            }
        }
    )
    val promptInfo = BiometricPrompt.PromptInfo.Builder()
        .setTitle(context.getString(R.string.login_biometria_titulo))
        .setSubtitle(context.getString(R.string.login_biometria_subtitulo))
        .setNegativeButtonText(context.getString(R.string.cancelar))
        .setAllowedAuthenticators(
            BiometricManager.Authenticators.BIOMETRIC_STRONG or
            BiometricManager.Authenticators.BIOMETRIC_WEAK
        )
        .build()
    prompt.authenticate(promptInfo)
}
''', "ui/screens/login/LoginScreen.kt")

add_code(r'''
suspend fun loginConBiometria(): AuthResult {
    val uid = session.userId.first()
    if (uid == SessionManager.NO_SESSION || auth.currentUser == null) {
        return AuthResult.Error("Iniciá sesión con email y contraseña primero")
    }
    return try {
        val usuario = getUsuario(uid)
            ?: return AuthResult.Error("No se encontró la cuenta. Iniciá sesión manualmente")
        session.guardarSesion(uid, usuario.email, usuario.nombre)
        AuthResult.Exito(usuario)
    } catch (e: Exception) {
        AuthResult.Error("No se pudo iniciar sesión: ${e.message}")
    }
}
''', "data/repository/AuthRepository.kt")

doc.add_page_break()

# --- Tema 3: Base de datos Firestore ----------------------------------------
doc.add_heading("3. Base de datos para compras y productos (Firestore)", level=1)
add_tag("OBLIGATORIO")
add_speaker("Bossio")

add_subhead("Que decir")
doc.add_paragraph(
    "Ahora el corazon de la app: la base de datos donde se guardan las compras y sus "
    "productos. Cada vez que un usuario registra una compra (a mano o con la foto del "
    "ticket), se crea un documento en la coleccion 'compras' de Cloud Firestore con el "
    "supermercado, la fecha, la hora, el total y la lista de productos con su cantidad "
    "y subtotal."
)
doc.add_paragraph(
    "Como dijimos en la introduccion, Firestore en Android no es 'solo nube': el SDK "
    "guarda automaticamente una copia local de esos documentos en una base de datos "
    "SQLite interna del telefono. Cuando guardamos o leemos una compra, primero se "
    "escribe/lee en esa base local, y recien despues se sincroniza con el servidor "
    "cuando hay conexion. Esto significa que la app funciona igual sin internet, y es "
    "la parte que cubre el requisito de 'base de datos local' de esta entrega."
)

add_subhead("Como lo hicimos")
doc.add_paragraph(
    "Tenemos un CompraRepository que envuelve la coleccion 'compras' de Firestore. "
    "guardarCompra() es una funcion suspend que arma un Map con los datos de la compra "
    "y lo agrega con compras.add(mapa).await() - el .await() viene de la libreria "
    "kotlinx-coroutines-play-services y convierte la Task de Firebase en una llamada "
    "de corrutina normal. Este metodo es el ejemplo concreto de 'Registrar compras en "
    "la Base de Datos'."
)

add_code(r'''
suspend fun guardarCompra(
    usuarioId    : String,
    fecha        : String,
    hora         : String,
    supermercado : String,
    total        : Double,
    productos    : List<Producto>,
    ticketUri    : String? = null
): String {
    val doc = compras.add(
        mapOf(
            "usuarioId"      to usuarioId,
            "fecha"          to fecha,
            "hora"           to hora,
            "supermercado"   to supermercado,
            "total"          to total,
            "ticketImageUri" to ticketUri,
            "productos"      to productos.map { it.toMap() }
        )
    ).await()
    return doc.id
}
''', "data/repository/CompraRepository.kt")

doc.add_page_break()

# --- Tema 4: Sincronizacion en la nube --------------------------------------
doc.add_heading("4. Sincronizacion en la nube", level=1)
add_tag("EXTRA")
add_speaker("Bossio")

add_subhead("Que decir")
doc.add_paragraph(
    "El segundo agregado diferencial es la sincronizacion en la nube en tiempo real. "
    "Si el usuario registra una compra desde el celular y despues abre la app en otro "
    "dispositivo (o la consola de Firebase), va a ver esa compra reflejada "
    "automaticamente, sin tener que refrescar nada."
)
doc.add_paragraph(
    "Esto lo logramos usando un 'listener' de Firestore: en lugar de pedir los datos "
    "una sola vez, nos suscribimos a la coleccion y Firestore nos avisa cada vez que "
    "algo cambia, tanto si el cambio vino de la nube como si vino de la copia local. "
    "Este mismo mecanismo es el que usamos para 'Listar compras guardadas'."
)

add_subhead("Como lo hicimos")
doc.add_paragraph(
    "getComprasFlow() esta implementado con callbackFlow, un constructor de Flow "
    "pensado para envolver callbacks. Dentro, llamamos a "
    "compras.whereEqualTo(\"usuarioId\", usuarioId).addSnapshotListener { snapshot, "
    "error -> ... }. Cada vez que el listener se dispara, mapeamos los documentos a "
    "objetos Compra, los ordenamos por fecha y los emitimos con trySend(lista). Cuando "
    "el Flow se cancela (por ejemplo al salir de la pantalla), awaitClose elimina el "
    "listener para no dejar fugas de memoria. El resultado: cualquier ViewModel que "
    "haga repo.getComprasFlow(userId).collect{} recibe la lista actualizada en tiempo "
    "real, tanto si el cambio vino del propio dispositivo como de otro."
)

add_code(r'''
fun getComprasFlow(usuarioId: String): Flow<List<Compra>> = callbackFlow {
    val registro = compras
        .whereEqualTo("usuarioId", usuarioId)
        .addSnapshotListener { snapshot, error ->
            if (error != null) {
                close(error)
                return@addSnapshotListener
            }
            val lista = snapshot?.documents
                ?.mapNotNull { it.toCompra() }
                ?.sortedByDescending { it.fecha.atTime(it.hora) }
                ?: emptyList()
            trySend(lista)
        }
    awaitClose { registro.remove() }
}
''', "data/repository/CompraRepository.kt")

doc.add_page_break()

# --- Tema 5: Corrutinas -----------------------------------------------------
doc.add_heading("5. Operaciones con corrutinas", level=1)
add_tag("OBLIGATORIO")
add_speaker("Bossio")

add_subhead("Que decir")
doc.add_paragraph(
    "Otro requisito obligatorio es usar corrutinas para las operaciones que tardan o "
    "dependen de I/O: leer/escribir en Firestore, llamar a una API, etc. Las "
    "corrutinas de Kotlin nos permiten escribir codigo asincrono que se lee como "
    "codigo secuencial, sin bloquear el hilo principal (el que dibuja la interfaz)."
)
doc.add_paragraph(
    "En toda la app, los ViewModel lanzan corrutinas con viewModelScope.launch, que "
    "automaticamente se cancelan si el usuario sale de la pantalla. Dentro de esas "
    "corrutinas llamamos a funciones suspend como las de SessionManager, "
    "CompraRepository o las APIs externas."
)

add_subhead("Como lo hicimos")
doc.add_paragraph(
    "Un ejemplo simple es ListadoComprasViewModel.cargar(): lanza una corrutina con "
    "viewModelScope.launch, primero obtiene el userId actual con "
    "session.userId.first() (espera el primer valor del Flow de DataStore), y si hay "
    "sesion, hace repo.getComprasFlow(usuarioId).collect { compras -> ... } para ir "
    "actualizando el estado de la pantalla cada vez que cambian las compras. Si se "
    "vuelve a llamar a cargar() (por ejemplo al refrescar), cancelamos el job anterior "
    "con observarJob?.cancel() antes de lanzar uno nuevo, para no tener dos colectores "
    "corriendo a la vez."
)

add_code(r'''
fun cargar() {
    observarJob?.cancel()
    observarJob = viewModelScope.launch {
        val usuarioId = session.userId.first()
        if (usuarioId == SessionManager.NO_SESSION) {
            _uiState.value = UiState.Error("Sesión expirada")
            return@launch
        }
        repo.getComprasFlow(usuarioId).collect { compras ->
            _uiState.value = UiState.Success(compras)
        }
    }
}
''', "ui/screens/compras/listado/ListadoComprasViewModel.kt")

doc.add_page_break()

# --- Tema 6: Menus y dialogos -----------------------------------------------
doc.add_heading("6. Menus y dialogos", level=1)
add_tag("OBLIGATORIO")
add_speaker("Bossio")

add_subhead("Que decir")
doc.add_paragraph(
    "El requisito de menus y dialogos lo cubrimos de varias formas. Por un lado, "
    "usamos AlertDialog de Jetpack Compose para confirmaciones destructivas: antes de "
    "borrar una compra del historial, mostramos un dialogo de confirmacion con un "
    "icono de advertencia, el texto '¿Seguro que querés eliminar esta compra?' y dos "
    "botones, Eliminar y Cancelar."
)
doc.add_paragraph(
    "Por otro lado, usamos menus desplegables (ExposedDropdownMenu y DropdownMenu) en "
    "varios lugares: por ejemplo, en el comparador de precios, para elegir que "
    "producto querer comparar entre supermercados aparece un menu desplegable con "
    "todos los productos detectados ese mes."
)

add_subhead("Como lo hicimos")
doc.add_paragraph(
    "El dialogo de confirmacion se maneja con una variable de estado "
    "compraAEliminar: Compra? - cuando no es null, se muestra un AlertDialog cuyo "
    "confirmButton llama a vm.eliminar(compra.id) y cierra el dialogo, y cuyo "
    "dismissButton solo lo cierra. Para el menu desplegable usamos "
    "ExposedDropdownMenuBox, que combina un OutlinedTextField de solo lectura con un "
    "ExposedDropdownMenu que lista las opciones; al tocar una opcion se llama a "
    "onSeleccionar(nombre) y el menu se cierra."
)

add_code(r'''
compraAEliminar?.let { compra ->
    AlertDialog(
        onDismissRequest = { compraAEliminar = null },
        icon = {
            Icon(Icons.Outlined.Warning, contentDescription = null,
                tint = MaterialTheme.colorScheme.error)
        },
        title = { Text(stringResource(R.string.historial_eliminar_titulo)) },
        text  = {
            Text(stringResource(R.string.listado_eliminar_mensaje,
                compra.supermercado,
                compra.fecha.format(DateTimeFormatter.ofPattern("dd/MM/yyyy"))))
        },
        confirmButton = {
            Button(
                onClick = { vm.eliminar(compra.id); compraAEliminar = null },
                colors  = ButtonDefaults.buttonColors(containerColor = MaterialTheme.colorScheme.error)
            ) { Text(stringResource(R.string.eliminar)) }
        },
        dismissButton = {
            OutlinedButton(onClick = { compraAEliminar = null }) {
                Text(stringResource(R.string.cancelar))
            }
        }
    )
}
''', "ui/screens/compras/listado/ListadoComprasScreen.kt")

add_code(r'''
@Composable
private fun ProductoSelector(
    productos     : List<String>,
    seleccionado  : String?,
    onSeleccionar : (String) -> Unit
) {
    var expanded by remember { mutableStateOf(false) }
    ExposedDropdownMenuBox(expanded = expanded, onExpandedChange = { expanded = it }) {
        OutlinedTextField(
            value = seleccionado ?: "", onValueChange = {}, readOnly = true,
            modifier = Modifier.fillMaxWidth().menuAnchor(),
            trailingIcon = { ExposedDropdownMenuDefaults.TrailingIcon(expanded) }
        )
        ExposedDropdownMenu(expanded = expanded, onDismissRequest = { expanded = false }) {
            productos.forEach { nombre ->
                DropdownMenuItem(
                    text    = { Text(nombre) },
                    onClick = { onSeleccionar(nombre); expanded = false }
                )
            }
        }
    }
}
''', "ui/screens/comparativa/ComparativaPreciosScreen.kt")

doc.add_page_break()

# ===========================================================================
# PARTE 2 - CORREA
# ===========================================================================
p = doc.add_paragraph()
r = p.add_run("PARTE 2 - CORREA")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = ACCENT
doc.add_paragraph(
    "Intents, networking con APIs publicas, carga de datos, IA, OCR, chat y "
    "comparativa de precios."
)

# --- Tema 7: Intents ---------------------------------------------------------
doc.add_heading("7. Compartir compra por Intent", level=1)
add_tag("OBLIGATORIO")
add_speaker("Correa")

add_subhead("Que decir")
doc.add_paragraph(
    "Ahora me toca a mi. Empiezo con el requisito de Intents. En la pantalla de "
    "detalle de una compra agregamos un boton de compartir (el icono de Share) que "
    "arma un texto con el resumen de la compra -supermercado, fecha, cada producto "
    "con su cantidad y precio, y el total- y lo manda a cualquier app que el usuario "
    "tenga instalada para compartir texto: WhatsApp, email, notas, lo que sea."
)

add_subhead("Como lo hicimos")
doc.add_paragraph(
    "Construimos el texto con buildString y appendLine, recorriendo la lista de "
    "productos de la compra. Despues creamos un Intent con la accion ACTION_SEND, le "
    "ponemos type = \"text/plain\" y le pasamos el texto en EXTRA_TEXT. Por ultimo, en "
    "vez de lanzar el Intent directo, lo envolvemos con "
    "Intent.createChooser(intent, \"Compartir compra\") y llamamos a "
    "context.startActivity(...) - eso hace que Android muestre el selector nativo de "
    "apps para compartir, en lugar de abrir siempre la misma app."
)

add_code(r'''
IconButton(onClick = {
    compra?.let { c ->
        val fmt   = DateTimeFormatter.ofPattern("dd/MM/yyyy")
        val texto = buildString {
            appendLine("🛒 Compra en ${c.supermercado}")
            appendLine("📅 ${c.fecha.format(fmt)} · ${c.hora}")
            appendLine()
            c.productos.forEach { p ->
                appendLine("• ${p.nombre}  ×${p.cantidad}  ${"$%,.2f".format(p.subtotal)}")
            }
            appendLine()
            appendLine("Total: ${"$%,.2f".format(c.total)}")
            appendLine()
            appendLine("Enviado desde Super Ahorro 💚")
        }
        val intent = Intent(Intent.ACTION_SEND).apply {
            type = "text/plain"
            putExtra(Intent.EXTRA_TEXT, texto)
        }
        context.startActivity(Intent.createChooser(intent, "Compartir compra"))
    }
}, enabled = compra != null) {
    Icon(Icons.Default.Share, null, tint = MaterialTheme.colorScheme.primary)
}
''', "ui/screens/compras/detalle/DetalleCompraScreen.kt")

doc.add_page_break()

# --- Tema 8: Networking API publica -------------------------------------------
doc.add_heading("8. Networking: API publica de supermercados", level=1)
add_tag("OBLIGATORIO")
add_speaker("Correa")

add_subhead("Que decir")
doc.add_paragraph(
    "El siguiente requisito obligatorio es networking: consumir una API real desde la "
    "app. Implementamos PromocionesRepository, que consulta las APIs publicas de "
    "catalogo de Carrefour Argentina y Chango Mas (ambas usan la plataforma VTEX, que "
    "expone un endpoint publico de busqueda de productos en formato JSON)."
)
doc.add_paragraph(
    "Con esos datos armamos la pantalla de Promociones, que muestra ofertas reales con "
    "su precio normal, precio con descuento y el porcentaje de ahorro, por "
    "supermercado. Este es el ejemplo concreto de 'Consultar supermercados o "
    "promociones desde API'."
)

add_subhead("Como lo hicimos")
doc.add_paragraph(
    "Usamos OkHttp como cliente HTTP, con timeouts configurados de conexion y "
    "lectura. Para cada supermercado armamos una URL con HttpUrl.Builder, agregando "
    "parametros de paginacion (_from / _to) para traer de a 50 productos, en dos "
    "paginas. obtenerProductos() hace el request con "
    "client.newCall(request).execute(), lee el body como String y lo parsea como un "
    "JSONArray. Todo esto corre dentro de withContext(Dispatchers.IO), porque son "
    "llamadas de red bloqueantes que no deben correr en el hilo principal."
)

add_code(r'''
class PromocionesRepository {
    private val client = OkHttpClient.Builder()
        .connectTimeout(20, TimeUnit.SECONDS)
        .readTimeout(30, TimeUnit.SECONDS)
        .build()

    private val fuentes = listOf(
        FuenteVtex("Carrefour",  "https://www.carrefour.com.ar"),
        FuenteVtex("Chango Más", "https://www.masonline.com.ar"),
    )

    suspend fun buscarPromocionesArgentina(): PromocionesResult = withContext(Dispatchers.IO) {
        val promociones = mutableListOf<Promocion>()
        for (fuente in fuentes) {
            try {
                for (pagina in 0 until 2) {
                    val from = pagina * 50
                    val to   = from + 49
                    val url = "${fuente.baseUrl}/api/catalog_system/pub/products/search"
                        .toHttpUrl().newBuilder()
                        .addQueryParameter("_from", from.toString())
                        .addQueryParameter("_to", to.toString())
                        .build()
                    obtenerProductos(url).forEach { producto ->
                        parsearPromocion(producto, fuente.supermercado)?.let(promociones::add)
                    }
                }
            } catch (e: Exception) { }
        }
        if (promociones.isEmpty()) PromocionesResult.Error("No se pudieron obtener promociones en este momento.")
        else PromocionesResult.Exito(promociones)
    }
''', "data/repository/PromocionesRepository.kt")

add_code(r'''
    private fun obtenerProductos(url: HttpUrl): List<JSONObject> {
        val request = Request.Builder()
            .url(url)
            .header("Accept", "application/json")
            .header("User-Agent", "Mozilla/5.0 (Android) SuperAhorroApp")
            .get()
            .build()

        val response = client.newCall(request).execute()
        val body      = response.body?.string()
        if (!response.isSuccessful || body == null) return emptyList()

        val items = JSONArray(body)
        return (0 until items.length()).map { items.getJSONObject(it) }
    }
''', "data/repository/PromocionesRepository.kt")

doc.add_page_break()

# --- Tema 9: Carga real de datos + notificaciones ----------------------------
doc.add_heading("9. Carga real de datos y notificaciones de ofertas", level=1)
add_tag("OBLIGATORIO + EXTRA")
add_speaker("Correa")

add_subhead("Que decir")
doc.add_paragraph(
    "Este punto combina dos cosas: el requisito obligatorio de 'carga real de datos' "
    "y el extra de 'notificaciones'. En la pantalla principal (Home), apenas el "
    "usuario entra, el HomeViewModel dispara dos cargas reales en paralelo: por un "
    "lado, escucha en tiempo real las compras del usuario desde Firestore para "
    "calcular el gasto del mes, la ultima compra y el grafico de gastos por dia. Por "
    "otro lado, llama a PromocionesRepository para traer las mejores ofertas vigentes "
    "de los supermercados."
)
doc.add_paragraph(
    "Esas ofertas no solo se muestran en la pantalla de Promociones: tambien "
    "alimentan una campanita de notificaciones arriba a la derecha del Home. Si hay "
    "ofertas nuevas que el usuario todavia no vio, aparece un punto rojo sobre la "
    "campana; al tocarla, se abre un menu desplegable con las 5 mejores promociones "
    "del momento."
)

add_subhead("Como lo hicimos")
doc.add_paragraph(
    "En el init del HomeViewModel llamamos a cargar() (que arma el resumen con datos "
    "de Firestore) y a cargarNotificaciones(). Esta ultima hace "
    "viewModelScope.launch { promoRepo.buscarPromocionesArgentina() }, y si el "
    "resultado es Exito, ordena las promociones por porcentaje de descuento "
    "(precio/precioSinDescuento) y se queda con las 5 mejores en un StateFlow "
    "_notificaciones. En HomeScreen, un DropdownMenu cuyo expanded esta atado a "
    "showNotificaciones muestra esa lista; el punto rojo se muestra con "
    "notificaciones.isNotEmpty() && !notificacionesVistas, y se oculta apenas el "
    "usuario abre el menu."
)

add_code(r'''
private fun cargarNotificaciones() {
    viewModelScope.launch {
        val resultado = promoRepo.buscarPromocionesArgentina()
        if (resultado is PromocionesResult.Exito) {
            // Ordena por mayor descuento (menor proporción precio/precioSinDescuento) primero.
            _notificaciones.value = resultado.promociones
                .sortedBy { promo ->
                    val original = promo.precioSinDescuento
                    if (original != null && original > 0) (promo.precio / original) else 1.0
                }
                .take(5)
        }
    }
}
''', "ui/screens/home/HomeViewModel.kt")

add_code(r'''
IconButton(onClick = {
    showNotificaciones = true
    notificacionesVistas = true
}) {
    Icon(Icons.Default.Notifications, null, tint = MaterialTheme.colorScheme.onSurfaceVariant)
}
if (notificaciones.isNotEmpty() && !notificacionesVistas) {
    Box(modifier = Modifier.size(8.dp).align(Alignment.TopEnd)
        .offset(x = (-10).dp, y = 10.dp).clip(CircleShape)
        .background(MaterialTheme.colorScheme.error))
}
DropdownMenu(expanded = showNotificaciones, onDismissRequest = { showNotificaciones = false }) {
    if (notificaciones.isEmpty()) {
        Text(stringResource(R.string.notificaciones_vacio))
    } else {
        notificaciones.forEach { promo ->
            DropdownMenuItem(text = { /* nombre, precio y descuento de la promo */ },
                onClick = { showNotificaciones = false; onPromocionesClick() })
        }
    }
}
''', "ui/screens/home/HomeScreen.kt")

doc.add_page_break()

# --- Tema 10: Carga automatica y OCR con IA ----------------------------------
doc.add_heading("10. Carga automatica y OCR del ticket con IA", level=1)
add_tag("EXTRA")
add_speaker("Correa")

add_subhead("Que decir")
doc.add_paragraph(
    "Uno de los agregados mas vistosos: cuando el usuario va a registrar una compra "
    "nueva, puede sacarle una foto al ticket del supermercado y la app completa "
    "automaticamente el supermercado, la fecha, la hora, el total y la lista de "
    "productos con sus precios. Esto cubre tanto 'carga automatica del ticket con IA' "
    "como 'OCR del ticket', porque en realidad es la misma funcionalidad: usamos un "
    "modelo de IA con vision (no un OCR tradicional separado) que lee la imagen "
    "directamente y devuelve los datos estructurados."
)

add_subhead("Como lo hicimos")
doc.add_paragraph(
    "Usamos la API de Groq (un proveedor de modelos de IA con API compatible con "
    "OpenAI) a traves de OkHttp. GroqRepository.analizarTicket(bitmap) convierte la "
    "foto a base64 y la manda junto con un prompt muy detallado que le explica a la "
    "IA los distintos formatos de tickets argentinos (almacenes, Carrefour, Coto, "
    "etc.) y le pide que devuelva SOLO un JSON con el supermercado, fecha, hora, total "
    "y la lista de productos. Todo esto corre con withContext(Dispatchers.IO) porque "
    "es una llamada de red."
)
doc.add_paragraph(
    "Desde la pantalla de Nueva Compra, NuevaCompraViewModel.analizarTicketConIA() "
    "toma la foto guardada, llama a groqRepo.analizarTicket(bitmap) dentro de un "
    "viewModelScope.launch, y si el resultado es Exito, reemplaza la lista de "
    "productos del formulario y completa los demas campos (supermercado, fecha, hora, "
    "total) a traves de un callback onResultado."
)

add_code(r'''
suspend fun analizarTicket(bitmap: Bitmap): GroqResult = withContext(Dispatchers.IO) {
    if (apiKey.isBlank()) {
        return@withContext GroqResult.Error("Falta configurar GROQ_API_KEY en local.properties")
    }
    try {
        val base64 = bitmapToBase64(bitmap)

        val prompt = """
Sos un experto en leer tickets de supermercados y almacenes de Argentina.
TAREA: Extraer todos los productos del ticket. Responder SOLO con JSON válido.

FORMATO A (almacenes, kioscos):
  1,0000 (00) x 4500,0000      ← línea de cantidad/precio unitario → IGNORAR
  COCA COLA RETORNABLE 4500,00 ← producto real (nombre + precio total)

FORMATO B (Carrefour, Jumbo, Coto, Disco):
  COCA COLA 1.5L
  1   $4.500,00   $4.500,00    ← cant, precio unit, precio total → tomar el último
... (el prompt sigue detallando los formatos y la estructura JSON esperada)
"""
        // POST a la API de Groq (modelo de visión) con el prompt + imagen en base64
        // -> parsea la respuesta JSON a un objeto Ticket(supermercado, fecha, hora, total, productos)
    } catch (e: Exception) {
        GroqResult.Error("No se pudo analizar el ticket: ${e.message}")
    }
}
''', "data/repository/GroqRepository.kt")

add_code(r'''
fun analizarTicketConIA(
    onResultado: (supermercado: String, fecha: String, hora: String, total: String, productos: List<Producto>) -> Unit
) {
    val uri = _ticketUri.value ?: run {
        _errorIA.value = "Primero tomá o seleccioná una foto del ticket"
        return
    }
    viewModelScope.launch {
        _analizando.value = true
        _errorIA.value    = null
        try {
            val bitmap = /* decodificar el bitmap a partir de la uri de la foto */
            when (val resultado = groqRepo.analizarTicket(bitmap)) {
                is GroqResult.Exito -> {
                    val t = resultado.ticket
                    _productos.value = t.productos
                    onResultado(t.supermercado, fechaIaValida(t.fecha), t.hora, t.total, t.productos)
                }
                is GroqResult.Error -> _errorIA.value = resultado.mensaje
            }
        } catch (e: Exception) {
            _errorIA.value = "No se pudo procesar la imagen: ${e.message}"
        } finally {
            _analizando.value = false
        }
    }
}
''', "ui/screens/compras/nueva/NuevaCompraViewModel.kt")

doc.add_page_break()

# --- Tema 11: Chat sobre historial -------------------------------------------
doc.add_heading("11. Chat de consultas sobre el historial", level=1)
add_tag("EXTRA")
add_speaker("Correa")

add_subhead("Que decir")
doc.add_paragraph(
    "Otro diferencial es un asistente de chat con IA al que el usuario le puede "
    "preguntar cosas sobre su propio historial de compras: '¿cuanto gaste el mes "
    "pasado en Carrefour?', '¿que es lo que mas compro?', etc. La IA responde en base "
    "a los datos reales del usuario, no con informacion generica."
)

add_subhead("Como lo hicimos")
doc.add_paragraph(
    "Cuando el usuario envia una pregunta, ChatViewModel.enviar() arma primero un "
    "'contexto': lee las compras del usuario con repo.getComprasFlow(userId).first() "
    "y genera un resumen en texto plano con las ultimas 40 compras (fecha, "
    "supermercado, total y cada producto con cantidad y subtotal). Despues manda ese "
    "contexto, la pregunta y los ultimos mensajes del chat (para mantener el hilo de "
    "la conversacion) a groqRepo.consultarHistorial(), que hace la llamada HTTP a la "
    "API de Groq con un modelo de texto. La respuesta se agrega a la lista de "
    "mensajes (_mensajes), que la pantalla de Chat muestra como burbujas."
)

add_code(r'''
fun enviar(pregunta: String) {
    val texto = pregunta.trim()
    if (texto.isBlank() || _escribiendo.value) return

    _mensajes.value = _mensajes.value + MensajeChat(texto, esUsuario = true)

    viewModelScope.launch {
        _escribiendo.value = true
        val contexto = armarContextoCompras()
        val previos = _mensajes.value.dropLast(1)
            .filterNot { it.esError }
            .takeLast(10)
            .map { (if (it.esUsuario) "user" else "assistant") to it.texto }

        val mensaje = when (val resultado = groqRepo.consultarHistorial(texto, contexto, previos)) {
            is ChatResult.Exito -> MensajeChat(resultado.respuesta, esUsuario = false)
            is ChatResult.Error -> MensajeChat(resultado.mensaje, esUsuario = false, esError = true)
        }
        _mensajes.value = _mensajes.value + mensaje
        _escribiendo.value = false
    }
}

private suspend fun armarContextoCompras(): String {
    val userId = session.userId.first()
    val compras = repo.getComprasFlow(userId).first()
    if (compras.isEmpty()) return "El usuario no tiene compras registradas."

    val fmt = DateTimeFormatter.ofPattern("dd/MM/yyyy")
    return buildString {
        appendLine("El usuario tiene ${compras.size} compras registradas:")
        compras.take(40).forEach { c ->
            appendLine("- ${c.fecha.format(fmt)} ${c.hora} | ${c.supermercado} | Total: $${"%.2f".format(c.total)}")
            c.productos.forEach { p ->
                appendLine("    * ${p.nombre} x${p.cantidad} = $${"%.2f".format(p.subtotal)}")
            }
        }
    }
}
''', "ui/screens/chat/ChatViewModel.kt")

doc.add_page_break()

# --- Tema 12: Comparativa de precios ------------------------------------------
doc.add_heading("12. Comparativa de precios entre supermercados", level=1)
add_tag("EXTRA")
add_speaker("Correa")

add_subhead("Que decir")
doc.add_paragraph(
    "El ultimo agregado diferencial es el comparador de precios. La idea es: de todas "
    "las compras que el usuario registro en un mes, agrupar los productos que son 'el "
    "mismo' aunque se llamen distinto en cada ticket (por ejemplo, 'COCA COLA 1.5L' en "
    "un supermercado y 'GASEOSA COCA COLA 1.5LT' en otro), y mostrar en que "
    "supermercado salio mas barato cada producto, marcando el mas economico con una "
    "estrella."
)

add_subhead("Como lo hicimos")
doc.add_paragraph(
    "ComparativaPreciosViewModel primero pide a CompraRepository los meses que tienen "
    "compras (getMesesConCompras) y, para el mes seleccionado, todos los productos "
    "comprados con su precio y supermercado (getProductosDelMes). Como esos nombres "
    "pueden venir escritos de formas distintas, se los pasamos a "
    "groqRepo.agruparProductos(), que usa un modelo de IA de texto para devolver "
    "grupos de nombres equivalentes (GrupoProductos). Con esos grupos armamos la lista "
    "de 'productos genericos' que se puede elegir en el menu desplegable del tema "
    "anterior, y comparacionDe(nombre) filtra y ordena por precio todas las "
    "apariciones de ese grupo en el mes, para mostrar cual es la mas barata."
)

add_code(r'''
private suspend fun cargarMes(meses: List<String>, mes: String) {
    _uiState.value = UiState.Success(/* ...estado de carga, cargandoIA = true */)

    productosDelMes = repo.getProductosDelMes(usuarioId, mes)
    val nombresDistintos = productosDelMes.map { it.nombre.trim() }.distinctBy { it.lowercase() }

    // La IA agrupa nombres equivalentes del mismo producto entre supermercados
    grupos = groqRepo.agruparProductos(nombresDistintos)
    val nombres = grupos.map { it.nombre }.distinctBy { it.lowercase() }.sortedBy { it.lowercase() }

    val productoSel = nombres.firstOrNull()
    _uiState.value = UiState.Success(
        ComparativaData(
            mesesDisponibles = meses, mesSeleccionado = mes,
            productosDisponibles = nombres, productoSeleccionado = productoSel,
            comparacion = productoSel?.let { comparacionDe(it) } ?: emptyList(),
            cargandoIA = false
        )
    )
}

private fun comparacionDe(nombreCanonico: String): List<PrecioProducto> {
    val itemsDelGrupo = grupos
        .filter { it.nombre.equals(nombreCanonico, ignoreCase = true) }
        .flatMap { it.items }
        .map { it.lowercase() }
        .toSet()

    return productosDelMes
        .filter { it.nombre.trim().lowercase() in itemsDelGrupo }
        .sortedBy { it.precio }
}
''', "ui/screens/comparativa/ComparativaPreciosViewModel.kt")

doc.add_page_break()

# ===========================================================================
# CIERRE (ambos)
# ===========================================================================
doc.add_heading("Cierre", level=1)
add_speaker("Bossio y Correa, en conjunto")

doc.add_paragraph(
    "Con esto cubrimos los siete requisitos obligatorios de la segunda entrega "
    "-persistencia de sesion con DataStore, base de datos de compras con Firestore "
    "(incluyendo su cache local), operaciones con corrutinas, networking contra una "
    "API publica real, menus y dialogos, carga real de datos e Intents- y seis de los "
    "agregados diferenciales de la etapa final: autenticacion biometrica, "
    "sincronizacion en la nube, notificaciones, carga automatica y OCR de tickets con "
    "IA, chat sobre el historial y comparador de precios."
)
doc.add_paragraph(
    "Como proximos pasos nos quedan pendientes dos extras que todavia no cerramos del "
    "todo: la exportacion de datos (CSV/PDF) y los filtros avanzados por fecha o "
    "categoria en el listado de compras -la busqueda por texto ya funciona, pero los "
    "filtros rapidos por chips son por ahora solo visuales."
)

doc.add_page_break()

# ===========================================================================
# ANEXO: Posibles preguntas del profesor
# ===========================================================================
doc.add_heading("Anexo: posibles preguntas del profesor", level=1)

preguntas = [
    ("¿Por que usan Firestore y no Room/SQLite si pedian base de datos local?",
     "Porque el SDK de Firestore para Android incluye, por defecto, una base de datos "
     "local SQLite que actua como cache: todas las operaciones (leer y escribir) se "
     "resuelven primero contra esa copia local, y se sincronizan con el servidor en "
     "segundo plano cuando hay conexion. Por eso la misma implementacion cumple el "
     "requisito de base de datos local de esta entrega y, a la vez, es la base de la "
     "sincronizacion en la nube que sumamos como diferencial."),
    ("¿Que pasa si no hay conexion a internet?",
     "Las pantallas siguen funcionando con los datos de la cache local de Firestore; "
     "cuando vuelve la conexion, los cambios pendientes se sincronizan "
     "automaticamente."),
    ("¿Donde usan corrutinas ademas de Firestore?",
     "En todas las llamadas a las APIs externas (Groq, VTEX/Promociones) con "
     "withContext(Dispatchers.IO), y en SessionManager con DataStore, que expone "
     "Flows que se consumen con .collect{} o .first() dentro de viewModelScope.launch."),
    ("¿La IA tiene acceso a internet / a una API key?",
     "Si, usamos la API de Groq con una API key configurada en local.properties (no se "
     "sube al repositorio por seguridad)."),
]

for q, a in preguntas:
    p = doc.add_paragraph()
    r = p.add_run("P: " + q)
    r.bold = True
    p2 = doc.add_paragraph()
    r2 = p2.add_run("R: " + a)
    doc.add_paragraph()

print("Documento completo generado")
doc.save("Guia_Exposicion_SuperAhorro.docx")
