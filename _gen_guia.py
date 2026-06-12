# -*- coding: utf-8 -*-
"""Genera la guía de presentación en Word (.docx)."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---------------------------------------------------------------- estilos base
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)

for i, (size, color) in enumerate([(22, '1F4E79'), (16, '1F4E79'), (13, '2E75B6')], start=1):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.color.rgb = RGBColor.from_string(color)
    h.font.bold = True

title_style = doc.styles['Title']
title_style.font.name = 'Calibri'
title_style.font.color.rgb = RGBColor.from_string('1F4E79')

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2)
section.right_margin = Cm(2)
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)


# ---------------------------------------------------------------- helpers
def shade_paragraph(paragraph, color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)


def add_code(doc, code, color='EDEDED'):
    lines = code.strip('\n').split('\n')
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Cm(0.4)
        shade_paragraph(p, color)
        run = p.add_run(line if line.strip() else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string('1F1F1F')
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_file_tag(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string('C0392B')


def add_note(doc, label, text):
    p = doc.add_paragraph()
    shade_paragraph(p, 'FFF2CC')
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f'{label}: ')
    run.bold = True
    p.add_run(text)


def add_warning(doc, label, text):
    p = doc.add_paragraph()
    shade_paragraph(p, 'F8D7DA')
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.color.rgb = RGBColor.from_string('842029')
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor.from_string('842029')


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p


def numbered(doc, text):
    return doc.add_paragraph(text, style='List Number')


def make_table(doc, headers, rows, widths=None, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    for row in table.rows:
        for cell in row.cells:
            for par in cell.paragraphs:
                for r in par.runs:
                    r.font.size = Pt(font_size)
    return table


# ================================================================== PORTADA
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('Guía de presentación')
run.font.size = Pt(30)
run.font.bold = True
run.font.color.rgb = RGBColor.from_string('1F4E79')

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = t2.add_run('SuperAhorro – App Android')
run2.font.size = Pt(20)
run2.font.color.rgb = RGBColor.from_string('2E75B6')

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = t3.add_run(
    'Segunda entrega + etapa final: app funcionando y casos de uso reales de\n'
    'persistencia local (DataStore), base de datos (Room), corrutinas, networking,\n'
    'menús/diálogos, Intents y agregado diferencial (IA, biometría, notificaciones, filtros)'
)
run3.font.size = Pt(12)
run3.italic = True

doc.add_paragraph()
p = doc.add_paragraph('Cómo usar esta guía:')
p.runs[0].bold = True
bullet(doc, 'Es tu guion personal: leela antes del viernes y tenela a mano (o impresa / en el celular) durante la presentación.')
bullet(doc, 'Cada sección tiene 3 partes: qué mostrar en la app funcionando, qué archivo/código abrir en Android Studio, y qué decir en voz alta.')
bullet(doc, 'No necesitás leer el código línea por línea: señalá las partes marcadas y explicá la idea con tus palabras.')
bullet(doc, 'La sección 2 es un mapa rápido: para cada punto de la consigna del profesor, te dice en qué parte de la app y del código está resuelto.')

doc.add_page_break()

# ================================================================== 0. CHECKLIST PREVIO
doc.add_heading('0. Antes de empezar – checklist', level=1)

bullet(doc, 'Probar el flujo completo el día anterior con el celular/emulador conectado a internet (Promociones y el análisis con IA necesitan red real).')
bullet(doc, 'Tener un usuario ya registrado de prueba (por si el registro en vivo falla) y otro para registrar en vivo.')
bullet(doc, 'Tener al menos 1-2 compras ya guardadas en el historial, además de la que vas a crear en vivo, para que listado/historial/estadísticas no se vean vacíos.')
bullet(doc, 'Sacarle (o tener guardada) una foto de un ticket real de supermercado para mostrar "Analizar con IA".')
bullet(doc, 'Verificar que GROQ_API_KEY esté configurada en local.properties (si no está, el análisis con IA va a fallar).')
bullet(doc, 'Si el celular tiene huella/face configurada, probar el login con biometría antes (necesita que el usuario haya iniciado sesión normal al menos una vez).')
bullet(doc, 'Abrir en Android Studio, en pestañas y en el orden de esta guía, los archivos que vas a mostrar.')
bullet(doc, 'Plan B: si se corta internet, mostrá el código igual, explicá qué haría, y mencioná que ya lo probaste antes (podés tener capturas de pantalla de respaldo).')

add_note(doc, 'Tip', 'Llevá esta guía como guion, pero no la leas textual frente al profesor: usala para no perder el hilo y para acordarte qué archivo va con cada parte.')

doc.add_page_break()

# ================================================================== 1. RESUMEN DEL PROYECTO
doc.add_heading('1. Resumen del proyecto y arquitectura (1-2 min)', level=1)

doc.add_paragraph(
    'SuperAhorro es una app Android escrita en Kotlin con Jetpack Compose, organizada con el patrón '
    'MVVM (Model-View-ViewModel). Permite registrar usuarios, guardar compras de supermercado con sus '
    'productos, ver historial y estadísticas, consultar promociones reales de supermercados, analizar '
    'tickets con IA, compartir compras y loguearse con biometría.'
)

p = doc.add_paragraph('Las capas de la app son siempre las mismas:')
add_code(doc, 'UI (Screen, Compose)  ->  ViewModel  ->  Repository  ->  Fuente de datos\n'
              '                                                     +-- DataStore (preferencias)\n'
              '                                                     +-- Room (base de datos SQLite)\n'
              '                                                     +-- OkHttp (APIs externas por internet)')

bullet(doc, 'La Screen (Compose) solo dibuja y observa un estado (uiState).', bold_prefix='UI: ')
bullet(doc, 'El ViewModel pide datos al Repository, corre corrutinas y expone el resultado como UiState (Loading / Success / Error).', bold_prefix='ViewModel: ')
bullet(doc, 'El Repository decide de dónde vienen los datos y los traduce a los modelos de la app.', bold_prefix='Repository: ')
bullet(doc, 'DataStore, Room y las llamadas a APIs externas (OkHttp).', bold_prefix='Fuente de datos: ')

add_note(doc, 'Qué decir', '"La app sigue siempre el mismo patrón: la pantalla le pide datos al ViewModel, el ViewModel '
               'le pide al Repository, y el Repository es el único que sabe si esos datos vienen de DataStore, '
               'de la base de datos local o de internet. Todas las operaciones que tardan (leer/escribir en disco, '
               'pedir datos por internet) corren en corrutinas, así la UI nunca se traba."')

doc.add_page_break()

# ================================================================== 2. MAPA DE LA CONSIGNA
doc.add_heading('2. Mapa de la consigna: qué pide y dónde está resuelto', level=1)

doc.add_paragraph(
    'Esta tabla es tu "chuleta" rápida: para cada ítem que pide el profesor, indica en qué pantalla de '
    'la app se ve y qué archivo de código mostrar. Cada fila se explica en detalle más adelante.'
)

doc.add_heading('Segunda entrega — requisitos obligatorios', level=2)
make_table(
    doc,
    ['Requisito de la consigna', 'Dónde se ve en la app', 'Archivo(s) clave', 'Sección'],
    [
        ('Persistencia local de sesión o preferencias',
         'Cerrar y reabrir la app sin perder el login; cambiar tema oscuro/claro',
         'SessionManager.kt\nThemeviewmodel.kt', 'A'),
        ('Base de datos local para compras y productos',
         'Registro/login, Nueva compra, Historial, Detalle de compra',
         'AppDatabase.kt, CompraEntity.kt,\nCompraDao.kt, CompraRepository.kt', 'B'),
        ('Operaciones con corrutinas',
         'Se usan en TODA la app (cada vez que se lee/escribe BD, DataStore o se llama a una API)',
         'viewModelScope.launch, withContext,\nFlow.collect, .first(), mapLatest', 'C'),
        ('Networking',
         'Pantalla "Promociones" (API real de supermercados) y "Analizar con IA" en Nueva compra',
         'PromocionesRepository.kt\nGroqRepository.kt', 'D'),
        ('Menús y diálogos',
         'Menú de notificaciones, menú de perfil, confirmaciones de borrado, selector de fecha, filtros',
         'HomeScreen.kt, NuevaCompraScreen.kt,\nHistorialComprasScreen.kt, etc.', 'E'),
        ('Carga real de datos',
         'Promociones (datos en vivo de Carrefour/Changomás); compras/productos reales en Room',
         'PromocionesRepository.kt\nCompraRepository.kt', 'D / B'),
        ('Al menos una funcionalidad con Intents',
         'Sacar/elegir foto del ticket (cámara o galería) y "Compartir compra" desde el detalle',
         'NuevaCompraScreen.kt\nDetalleCompraScreen.kt', 'F'),
    ],
    font_size=9
)

doc.add_paragraph()
doc.add_heading('Ejemplos concretos mencionados por la consigna', level=2)
make_table(
    doc,
    ['Ejemplo concreto', 'Cómo está implementado'],
    [
        ('Guardar usuario logueado',
         'AuthRepository.login()/registrar() valida/crea el usuario en Room y guarda su id en DataStore '
         '(SessionManager) para que la sesión persista.'),
        ('Registrar compras en la Base de Datos',
         'NuevaCompraViewModel.guardar() -> CompraRepository.guardarCompra() inserta la compra y sus '
         'productos en Room (tablas "compras" y "productos").'),
        ('Listar compras guardadas',
         'ListadoComprasViewModel / HistorialComprasViewModel leen CompraRepository.getComprasFlow(), '
         'un Flow de Room que actualiza la lista automáticamente.'),
        ('Consultar supermercados o promociones desde API',
         'PromocionesRepository consulta en vivo las APIs públicas (VTEX) de Carrefour y Changomás.'),
        ('Compartir una compra por Intent',
         'DetalleCompraScreen: botón "compartir" arma un Intent.ACTION_SEND con el detalle de la compra '
         'y abre el selector de apps con Intent.createChooser().'),
    ],
    font_size=9.5
)

doc.add_paragraph()
doc.add_heading('Etapa final — agregado diferencial', level=2)
make_table(
    doc,
    ['Opcional de la consigna', 'Estado en SuperAhorro'],
    [
        ('Carga automática del ticket con IA', 'Implementado (GroqRepository + NuevaCompraViewModel.analizarTicketConIA)'),
        ('OCR del ticket', 'Implementado como parte de lo anterior (la IA lee la foto del ticket)'),
        ('Autenticación biométrica', 'Implementado (LoginScreen: BiometricPrompt / huella o rostro)'),
        ('Notificaciones', 'Implementado como notificaciones in-app (HomeScreen: menú con ofertas nuevas, datos reales de la API)'),
        ('Filtros avanzados', 'Implementado: Historial filtra por año y mes (HistorialComprasViewModel); Listado filtra por búsqueda de texto'),
        ('Comparativa de precios entre supermercados', 'Parcial: Promociones muestra ofertas de 2 supermercados y Estadísticas muestra gasto por supermercado'),
        ('Exportación de datos (CSV/PDF)', 'Pendiente: existe la opción en Configuración pero todavía no tiene acción asociada'),
        ('Chat para consultas sobre historial de compras', 'No implementado'),
    ],
    font_size=9.5
)

add_note(doc, 'Qué decir', '"Cubrimos los 7 puntos obligatorios de la segunda entrega con funcionalidad real, y de la '
               'etapa final ya tenemos implementados 4 de los 8 opcionales: IA/OCR del ticket, biometría, '
               'notificaciones de ofertas y filtros avanzados. La comparativa entre supermercados está parcialmente '
               'cubierta por Promociones y Estadísticas, y exportación / chat quedan como próximos pasos."')

doc.add_page_break()

# ================================================================== 3. ORDEN SUGERIDO
doc.add_heading('3. Orden sugerido de la presentación', level=1)

make_table(
    doc,
    ['Paso', 'Qué mostrar', 'Requisito que cubre'],
    [
        ('1. Intro', 'App funcionando: splash -> login -> home -> navegación general', 'Contexto / arquitectura'),
        ('2. DataStore', 'Cerrar y volver a abrir la app sin perder la sesión; cambiar el tema oscuro/claro', 'Persistencia local (A)'),
        ('3. Room', 'Registrar usuario, login, crear una compra nueva con productos y verla en el historial', 'Base de datos local (B)'),
        ('4. Corrutinas', 'Se menciona "al pasar" mientras se muestra Room/Networking (no necesita demo aparte)', 'Corrutinas (C)'),
        ('5. Networking', 'Pantalla de Promociones (consume API pública de supermercados en vivo)', 'Networking / carga real (D)'),
        ('6. Menús y diálogos', 'Menú de notificaciones, menú de perfil, confirmación de borrado, selector de fecha/mes', 'Menús y diálogos (E)'),
        ('7. Intents', 'Nueva compra: elegir foto (cámara/galería) + analizar con IA; Detalle: compartir compra', 'Intents (F)'),
        ('8. Etapa final', 'Login con biometría; filtros por año/mes en Historial; notificaciones de ofertas reales', 'Agregado diferencial (G)'),
        ('9. Cierre', 'Repaso rápido de la arquitectura (capas) y preguntas', '-'),
    ],
    font_size=10
)

doc.add_paragraph()
add_note(doc, 'Por qué este orden', 'Vas mostrando la app funcionando y, inmediatamente después de cada demo, '
               'el código que la hace posible. Así el profesor ve primero el resultado y después el "cómo". '
               'Si el tiempo aprieta, los pasos 1 a 7 son los más importantes (cubren TODA la segunda entrega); '
               'el paso 8 es el "plus" que suma para la nota de etapa final.')

doc.add_page_break()

# ================================================================== A. DATASTORE
doc.add_heading('A. Persistencia local — sesión y preferencias (DataStore)', level=1)

doc.add_heading('¿Qué es y para qué se usa?', level=2)
doc.add_paragraph(
    'Jetpack DataStore Preferences es el reemplazo moderno de SharedPreferences: guarda pares clave-valor '
    'de forma asincrónica (con corrutinas) y reactiva (con Flow). En SuperAhorro se usa para:'
)
bullet(doc, 'Guardar la sesión del usuario logueado (id, email, nombre) para no pedir login cada vez que se abre la app.', bold_prefix='SessionManager: ')
bullet(doc, 'Guardar la preferencia de tema oscuro / claro / según el sistema.', bold_prefix='ThemeViewModel: ')

doc.add_heading('Demo en vivo', level=2)
numbered(doc, 'Iniciar sesión con un usuario (email + contraseña).')
numbered(doc, 'Cerrar la app completamente (deslizar desde "apps recientes", no solo apretar atrás).')
numbered(doc, 'Volver a abrir la app: mostrar que entra directo al Home sin pedir login de nuevo.')
numbered(doc, 'Ir a Configuración, cambiar el tema a oscuro, cerrar y reabrir la app: el tema oscuro persiste.')

doc.add_heading('Código para mostrar', level=2)

add_file_tag(doc, 'data/local/SessionManager.kt')
doc.add_paragraph('Es la clase que encapsula todo el acceso al DataStore de la sesión.')
add_code(doc, '''
private val Context.dataStore: DataStore<Preferences>
    by preferencesDataStore(name = "session")

class SessionManager(private val context: Context) {

    companion object {
        private val KEY_USER_ID    = intPreferencesKey("user_id")
        private val KEY_USER_EMAIL = stringPreferencesKey("user_email")
        private val KEY_USER_NAME  = stringPreferencesKey("user_name")
        const val NO_SESSION = -1
    }

    val userId: Flow<Int> = context.dataStore.data.map { prefs ->
        prefs[KEY_USER_ID] ?: NO_SESSION
    }

    suspend fun guardarSesion(userId: Int, email: String, nombre: String) {
        context.dataStore.edit { prefs ->
            prefs[KEY_USER_ID]    = userId
            prefs[KEY_USER_EMAIL] = email
            prefs[KEY_USER_NAME]  = nombre
        }
    }

    suspend fun cerrarSesion() {
        context.dataStore.edit { prefs -> prefs[KEY_USER_ID] = NO_SESSION /* ... */ }
    }
}
''')

bullet(doc, 'preferencesDataStore(name = "session") crea (una sola vez) un DataStore llamado "session".', bold_prefix='Línea clave 1: ')
bullet(doc, 'userId es un Flow: cualquier pantalla que lo observe se entera sola cuando cambia la sesión.', bold_prefix='Línea clave 2: ')
bullet(doc, 'dataStore.edit { } es la forma de escribir: es "suspend" porque DataStore siempre es asincrónico.', bold_prefix='Línea clave 3: ')

add_file_tag(doc, 'data/repository/AuthRepository.kt  ->  función login()')
doc.add_paragraph('Acá se conecta el login (que valida contra Room, sección B) con el guardado de sesión en DataStore.')
add_code(doc, '''
suspend fun login(email: String, password: String): AuthResult {
    val usuario = dao.login(email.trim().lowercase(), password)   // 1) valida en Room
    return if (usuario != null) {
        session.guardarSesion(usuario.id, usuario.email, usuario.nombre) // 2) guarda en DataStore
        AuthResult.Exito(usuario)
    } else {
        AuthResult.Error("Email o contraseña incorrectos")
    }
}
''')

add_file_tag(doc, 'ui/theme/Themeviewmodel.kt  (segundo DataStore independiente)')
doc.add_paragraph('Mismo mecanismo, pero con otro archivo de preferencias ("theme_prefs") para el tema oscuro/claro.')
add_code(doc, '''
private val Context.themeDataStore: DataStore<Preferences>
        by preferencesDataStore(name = "theme_prefs")

fun setDarkMode(enabled: Boolean) {
    _isDarkMode.value = enabled
    viewModelScope.launch {
        dataStore.edit { prefs ->
            prefs[KEY_USE_SYSTEM] = false
            prefs[KEY_IS_DARK]    = enabled
        }
    }
}
''')

add_note(doc, 'Qué decir', '"Uso DataStore para todo lo que es preferencias simples del usuario: la sesión activa y el '
               'tema. Es asincrónico, así que nunca bloquea la UI, y como devuelve un Flow, la pantalla se actualiza '
               'sola cuando el valor cambia. Por ejemplo, al loguearme se guarda el id de usuario en DataStore, y la '
               'próxima vez que abro la app, ese mismo id se usa para saltar directo al Home."')

doc.add_page_break()

# ================================================================== B. ROOM
doc.add_heading('B. Base de datos local — compras y productos (Room)', level=1)

doc.add_heading('¿Qué es y para qué se usa?', level=2)
doc.add_paragraph(
    'Room es la librería de Android para trabajar con SQLite de forma estructurada: se definen Entidades '
    '(tablas), DAOs (consultas) y una Database (configuración general). En SuperAhorro se usa para guardar '
    'usuarios y sus compras (con los productos de cada compra).'
)
bullet(doc, 'Tabla "usuarios": registro y login.', bold_prefix='UsuarioEntity: ')
bullet(doc, 'Tabla "compras": una compra hecha en un supermercado (fecha, hora, total, foto del ticket).', bold_prefix='CompraEntity: ')
bullet(doc, 'Tabla "productos": los ítems de cada compra (relación 1 a N con compras).', bold_prefix='ProductoEntity: ')

doc.add_heading('Demo en vivo', level=2)
numbered(doc, 'Registrar un usuario nuevo (pantalla Registro) -> queda guardado en la tabla "usuarios".')
numbered(doc, 'Cerrar sesión e iniciar sesión con ese usuario nuevo (valida contra la base local).')
numbered(doc, 'Ir a "Nueva compra": elegir supermercado, fecha/hora, agregar 1 o 2 productos a mano y guardar.')
numbered(doc, 'Ir a "Historial de compras": la compra recién creada aparece automáticamente en la lista (sin recargar nada).')
numbered(doc, 'Entrar al detalle de esa compra para mostrar los productos guardados.')
numbered(doc, '(Opcional) Eliminar un producto o la compra y mostrar que desaparece sola de la lista al instante.')

doc.add_heading('Código para mostrar', level=2)

add_file_tag(doc, 'data/local/AppDatabase.kt')
doc.add_paragraph('La configuración general de la base de datos: qué tablas existen y cómo se obtiene la instancia (singleton).')
add_code(doc, '''
@Database(
    entities = [UsuarioEntity::class, CompraEntity::class, ProductoEntity::class],
    version = 2,
    exportSchema = false
)
abstract class AppDatabase : RoomDatabase() {
    abstract fun usuarioDao(): UsuarioDao
    abstract fun compraDao(): CompraDao
    abstract fun productoDao(): ProductoDao

    companion object {
        @Volatile private var INSTANCE: AppDatabase? = null
        fun getInstance(context: Context): AppDatabase =
            INSTANCE ?: synchronized(this) {
                Room.databaseBuilder(context.applicationContext, AppDatabase::class.java, "superahorro.db")
                    .fallbackToDestructiveMigration()
                    .build()
                    .also { INSTANCE = it }
            }
    }
}
''')

add_file_tag(doc, 'data/local/CompraEntity.kt  (relación compras -> productos)')
add_code(doc, '''
@Entity(tableName = "compras")
data class CompraEntity(
    @PrimaryKey(autoGenerate = true) val id: Int = 0,
    val usuarioId: Int,
    val fecha: String, val hora: String,
    val supermercado: String, val total: Double,
    val ticketImageUri: String? = null
)

@Entity(
    tableName = "productos",
    foreignKeys = [ForeignKey(
        entity = CompraEntity::class, parentColumns = ["id"], childColumns = ["compraId"],
        onDelete = ForeignKey.CASCADE   // al borrar una compra, se borran sus productos
    )]
)
data class ProductoEntity(
    @PrimaryKey(autoGenerate = true) val id: Int = 0,
    val compraId: Int, val codigo: String, val nombre: String,
    val descripcion: String, val cantidad: Int, val precio: Double
)
''')

add_file_tag(doc, 'data/local/CompraDao.kt')
add_code(doc, '''
@Dao
interface CompraDao {
    @Insert(onConflict = OnConflictStrategy.REPLACE)
    suspend fun insertar(compra: CompraEntity): Long

    @Query("DELETE FROM compras WHERE id = :id")
    suspend fun eliminar(id: Int)

    // Flow: la UI se actualiza sola cuando cambian los datos
    @Query("SELECT * FROM compras WHERE usuarioId = :usuarioId ORDER BY fecha DESC, hora DESC")
    fun getComprasDeUsuario(usuarioId: Int): Flow<List<CompraEntity>>
}
''')

add_file_tag(doc, 'data/repository/CompraRepository.kt')
doc.add_paragraph('El Repository combina las compras con sus productos y devuelve el resultado como modelos de dominio.')
add_code(doc, '''
fun getComprasFlow(usuarioId: Int): Flow<List<Compra>> =
    compraDao.getComprasDeUsuario(usuarioId).mapLatest { entities ->
        entities.mapNotNull { entity ->
            val productos = productoDao.getProductosDeCompra(entity.id)
            entity.toDomain(productos)
        }
    }

suspend fun guardarCompra(usuarioId: Int, fecha: String, hora: String, supermercado: String,
                           total: Double, productos: List<Producto>, ticketUri: String? = null): Int {
    val compraId = compraDao.insertar(CompraEntity(usuarioId = usuarioId, fecha = fecha, hora = hora,
        supermercado = supermercado, total = total, ticketImageUri = ticketUri)).toInt()
    if (productos.isNotEmpty()) productoDao.insertarTodos(productos.map { it.toEntity(compraId) })
    return compraId
}
''')

add_note(doc, 'Qué decir', '"Para usuarios y compras uso Room. Cada compra tiene sus productos relacionados por '
               'una clave foránea con borrado en cascada: si borro la compra, se borran solos sus productos. '
               'El DAO devuelve un Flow, así que cuando guardo una compra nueva, la pantalla de historial se '
               'actualiza sola, sin recargar nada a mano. Esto también conecta con la sección A: para saber de '
               'qué usuario traer las compras, primero leo el id guardado en DataStore."')

doc.add_page_break()

# ================================================================== C. CORRUTINAS
doc.add_heading('C. Operaciones con corrutinas', level=1)

doc.add_heading('¿Qué son y para qué se usan?', level=2)
doc.add_paragraph(
    'Las corrutinas (kotlinx.coroutines) permiten ejecutar tareas que tardan (leer/escribir en disco, '
    'esperar una respuesta de internet) sin congelar la pantalla. En SuperAhorro aparecen en TODOS los '
    'ViewModels y Repositories, con 4 patrones que se repiten en toda la app:'
)
bullet(doc, 'Cada acción del usuario (login, guardar compra, cargar promociones) se ejecuta dentro de viewModelScope.launch, una corrutina ligada al ciclo de vida del ViewModel.', bold_prefix='viewModelScope.launch: ')
bullet(doc, 'Para tareas pesadas como llamadas HTTP, se fuerza explícitamente un hilo de fondo (no el principal/UI).', bold_prefix='withContext(Dispatchers.IO): ')
bullet(doc, 'Las listas que vienen de Room (compras, productos) son Flow: el ViewModel hace collect y la UI se redibuja sola cada vez que cambian los datos.', bold_prefix='Flow.collect: ')
bullet(doc, 'Para leer "una sola vez" un valor de un Flow (por ejemplo, el id de usuario guardado en DataStore antes de pedir sus compras).', bold_prefix='.first(): ')

doc.add_heading('Demo en vivo', level=2)
doc.add_paragraph(
    'No requiere una demo aparte: se "ve" en que la app nunca se traba mientras carga datos (Room, '
    'DataStore o Promociones) y en que las listas se actualizan solas. Se puede mencionar al pasar mientras '
    'se muestran las secciones B (Room) y D (Networking).'
)

doc.add_heading('Código para mostrar', level=2)

add_file_tag(doc, 'ui/screens/compras/listado/ListadoComprasViewModel.kt  (los 4 patrones juntos)')
add_code(doc, '''
fun cargar() {
    observarJob?.cancel()
    observarJob = viewModelScope.launch {              // 1) corrutina ligada al ViewModel
        val usuarioId = session.userId.first()         // 2) lee UNA VEZ el id desde DataStore
        repo.getComprasFlow(usuarioId).collect { compras -> // 3) escucha cambios en Room (Flow)
            _uiState.value = UiState.Success(compras)  // 4) actualiza el estado -> la UI se redibuja
        }
    }
}
''')

add_file_tag(doc, 'data/repository/CompraRepository.kt  (mapLatest)')
doc.add_paragraph('mapLatest transforma cada lista nueva que llega del Flow de Room (entidades) en modelos de dominio, sin bloquear nada.')
add_code(doc, '''
fun getComprasFlow(usuarioId: Int): Flow<List<Compra>> =
    compraDao.getComprasDeUsuario(usuarioId).mapLatest { entities ->
        entities.mapNotNull { entity ->
            val productos = productoDao.getProductosDeCompra(entity.id)  // suspend: espera en segundo plano
            entity.toDomain(productos)
        }
    }
''')

add_file_tag(doc, 'data/repository/PromocionesRepository.kt  (withContext(Dispatchers.IO))')
doc.add_paragraph('Las llamadas de red NUNCA deben correr en el hilo principal: por eso todo el bloque corre explícitamente en Dispatchers.IO.')
add_code(doc, '''
suspend fun buscarPromocionesArgentina(): PromocionesResult = withContext(Dispatchers.IO) {
    val promociones = mutableListOf<Promocion>()
    for (fuente in fuentes) {
        // ... requests HTTP con OkHttp (bloqueantes, pero corren en el hilo de IO) ...
    }
    if (promociones.isEmpty()) PromocionesResult.Error("No se pudieron obtener promociones")
    else PromocionesResult.Exito(promociones)
}
''')

add_note(doc, 'Qué decir', '"Toda operación que tarda -leer la base, leer DataStore o llamar a una API- se hace en '
               'una corrutina. Las pantallas que muestran listas (historial, listado, home) escuchan un Flow que '
               'viene de Room con collect, así que cuando guardo una compra nueva, esas pantallas se actualizan '
               'solas sin que yo tenga que recargar nada. Y las llamadas de red corren explícitamente en '
               'Dispatchers.IO para no bloquear la interfaz."')

doc.add_page_break()

# ================================================================== D. NETWORKING
doc.add_heading('D. Networking — APIs reales (carga real de datos)', level=1)

doc.add_heading('¿Qué es y para qué se usa?', level=2)
doc.add_paragraph(
    'La app usa OkHttp para hacer peticiones HTTP a APIs externas reales (no mockeadas) y parsea las '
    'respuestas JSON con org.json. Hay dos integraciones de este tipo en el proyecto:'
)
bullet(doc, 'Consulta los catálogos públicos (API VTEX) de Carrefour y Changomás para armar una lista de productos con descuento real. Esto también alimenta las notificaciones de "nueva oferta" (sección G).', bold_prefix='PromocionesRepository: ')
bullet(doc, 'Envía la foto de un ticket a la API de Groq (IA) y recibe en JSON los productos, fecha, hora y total del ticket. Es la base del agregado diferencial de IA/OCR (sección G).', bold_prefix='GroqRepository: ')

doc.add_heading('Demo en vivo', level=2)
numbered(doc, 'Entrar a la pantalla "Promociones": se ve un indicador de carga y después una lista de productos con precio rebajado, traídos en vivo de Carrefour y Changomás.')
numbered(doc, '(Opcional) Mostrar el filtro por ciudad/supermercado.')
numbered(doc, 'En "Nueva compra", tocar "Analizar ticket con IA" y elegir la foto de un ticket: la app llama a la API de Groq y completa automáticamente supermercado, fecha, hora, total y productos (ver también sección F y G).')

doc.add_heading('Código para mostrar', level=2)

add_file_tag(doc, 'data/repository/PromocionesRepository.kt')
doc.add_paragraph('Cliente HTTP, lista de fuentes (supermercados) y función principal que recorre cada una.')
add_code(doc, '''
class PromocionesRepository {
    private val client = OkHttpClient.Builder()
        .connectTimeout(20, TimeUnit.SECONDS)
        .readTimeout(30, TimeUnit.SECONDS)
        .build()

    private val fuentes = listOf(
        FuenteVtex("Carrefour",  "https://www.carrefour.com.ar"),
        FuenteVtex("Chango Más", "https://www.masonline.com.ar"),
    )

    suspend fun buscarPromocionesArgentina(): PromocionesResult = withContext(Dispatchers.IO) {
        val promociones = mutableListOf<Promocion>()
        for (fuente in fuentes) {
            try {
                for (pagina in 0 until 2) {
                    val url = (fuente.baseUrl + "/api/catalog_system/pub/products/search")
                        .toHttpUrl().newBuilder()
                        .addQueryParameter("_from", (pagina * 50).toString())
                        .addQueryParameter("_to", (pagina * 50 + 49).toString())
                        .build()
                    obtenerProductos(url).forEach { producto ->
                        parsearPromocion(producto, fuente.supermercado)?.let(promociones::add)
                    }
                }
            } catch (e: Exception) { /* si falla un super, sigue con el resto */ }
        }
        if (promociones.isEmpty()) PromocionesResult.Error("No se pudieron obtener promociones")
        else PromocionesResult.Exito(promociones)
    }
''')

add_file_tag(doc, 'data/repository/PromocionesRepository.kt  (continuación: request + parseo)')
add_code(doc, '''
    private fun obtenerProductos(url: HttpUrl): List<JSONObject> {
        val request = Request.Builder()
            .url(url)
            .header("Accept", "application/json")
            .header("User-Agent", "Mozilla/5.0 (Android) SuperAhorroApp")
            .get()
            .build()

        val response = client.newCall(request).execute()
        val body = response.body?.string()
        if (!response.isSuccessful || body == null) return emptyList()

        val items = JSONArray(body)
        return (0 until items.length()).map { items.getJSONObject(it) }
    }

    // Solo se considera "promoción" si el precio actual es menor al de lista
    private fun parsearPromocion(producto: JSONObject, supermercado: String): Promocion? {
        val offer = producto.optJSONArray("items")?.optJSONObject(0)
            ?.optJSONArray("sellers")?.optJSONObject(0)
            ?.optJSONObject("commertialOffer") ?: return null

        val precio = offer.optDouble("Price")
        val precioDeLista = offer.optDouble("ListPrice")
        if (precio.isNaN() || precioDeLista.isNaN() || precioDeLista <= precio) return null

        return Promocion(producto = producto.optString("productName"), precio = precio,
            precioSinDescuento = precioDeLista, moneda = "ARS",
            supermercado = supermercado, ciudad = "Córdoba", pais = "Argentina")
    }
}
''')

add_file_tag(doc, 'ui/screens/promociones/PromocionesViewModel.kt')
add_code(doc, '''
class PromocionesViewModel : ViewModel() {
    private val repo = PromocionesRepository()
    private val _uiState = MutableStateFlow<UiState<List<Promocion>>>(UiState.Loading)
    val uiState = _uiState.asStateFlow()

    init { cargar() }

    fun cargar() {
        viewModelScope.launch {
            _uiState.value = UiState.Loading
            _uiState.value = when (val resultado = repo.buscarPromocionesArgentina()) {
                is PromocionesResult.Exito -> UiState.Success(resultado.promociones)
                is PromocionesResult.Error -> UiState.Error(resultado.mensaje)
            }
        }
    }
}
''')

add_file_tag(doc, 'data/repository/GroqRepository.kt  (segunda API: IA para tickets)')
doc.add_paragraph('Mismo patrón (OkHttp + JSON), pero con un POST que manda la imagen del ticket en base64 a una IA.')
add_code(doc, '''
val request = Request.Builder()
    .url("https://api.groq.com/openai/v1/chat/completions")
    .addHeader("Authorization", "Bearer $apiKey")
    .addHeader("Content-Type", "application/json")
    .post(requestBody.toRequestBody("application/json".toMediaType()))
    .build()

val response = client.newCall(request).execute()
val body = response.body?.string() ?: return@withContext GroqResult.Error("Respuesta vacía de la API")
if (!response.isSuccessful) return@withContext GroqResult.Error("Error API ${response.code}: $body")

val content = JSONObject(body).getJSONArray("choices").getJSONObject(0)
    .getJSONObject("message").getString("content").trim()
// ... se parsea ese JSON a un objeto TicketAnalizado con los productos
''')

add_note(doc, 'Qué decir', '"Para promociones consumo directamente las APIs públicas que usan las webs de Carrefour y '
               'Changomás (son APIs VTEX, un estándar de e-commerce). Hago el request con OkHttp en un hilo de fondo '
               '(Dispatchers.IO), parseo el JSON con org.json y me quedo solo con los productos donde el precio actual '
               'es menor al precio de lista, es decir, los que tienen un descuento real. Si un supermercado falla, '
               'sigo con el resto, y si fallan todos devuelvo un estado de error que la UI muestra como tal. El '
               'análisis de tickets con IA usa el mismo patrón pero contra la API de Groq, mandando la foto en base64 '
               'y recibiendo un JSON con los productos del ticket: esto es carga real de datos en los dos sentidos, '
               'no hay nada mockeado."')

doc.add_page_break()

# ================================================================== E. MENUS Y DIALOGOS
doc.add_heading('E. Menús y diálogos', level=1)

doc.add_heading('¿Qué es y para qué se usa?', level=2)
doc.add_paragraph(
    'Jetpack Compose ofrece varios componentes para menús y ventanas emergentes. SuperAhorro usa varios '
    'tipos distintos, cada uno para un caso de uso concreto:'
)
bullet(doc, 'Lista desplegable con las últimas ofertas (notificaciones), en el ícono de campana del Home.', bold_prefix='DropdownMenu: ')
bullet(doc, 'Panel deslizable desde abajo para el menú de perfil (Mi Perfil, Mis Estadísticas, Cerrar Sesión) y para "agregar producto" en Nueva compra.', bold_prefix='ModalBottomSheet: ')
bullet(doc, 'Confirmaciones de acciones importantes: cerrar sesión, eliminar una compra, eliminar un producto.', bold_prefix='AlertDialog: ')
bullet(doc, 'Diálogos a medida para elegir el origen de la foto (cámara/galería) y para agregar un supermercado nuevo a la lista.', bold_prefix='Dialog (custom): ')
bullet(doc, 'Selector de fecha para la compra.', bold_prefix='DatePickerDialog: ')
bullet(doc, 'Filtros por mes (Historial) y por categoría/período (Listado, Estadísticas).', bold_prefix='FilterChip: ')

doc.add_heading('Demo en vivo', level=2)
numbered(doc, 'En el Home, tocar el ícono de la campana: se abre un DropdownMenu con las últimas ofertas (notificaciones).')
numbered(doc, 'Tocar el avatar de perfil en el Home: se abre un ModalBottomSheet con el menú (Mi Perfil, Estadísticas, Cerrar sesión).')
numbered(doc, 'Elegir "Cerrar sesión": aparece un AlertDialog de confirmación antes de cerrar la sesión.')
numbered(doc, 'En "Nueva compra": tocar el campo de fecha -> se abre el DatePickerDialog; tocar la cámara/galería -> diálogo a medida para elegir el origen de la foto.')
numbered(doc, 'En "Historial de compras": usar los chips de mes para filtrar las compras de ese mes.')
numbered(doc, 'En "Listado" o "Detalle de compra": eliminar un ítem y mostrar el AlertDialog de confirmación.')

doc.add_heading('Código para mostrar', level=2)

add_file_tag(doc, 'ui/screens/home/HomeScreen.kt  (DropdownMenu de notificaciones)')
doc.add_paragraph('El menú muestra las promociones (traídas por la sección D) como si fueran notificaciones de "nueva oferta".')
add_code(doc, '''
var showNotificaciones by remember { mutableStateOf(false) }

IconButton(onClick = { showNotificaciones = true }) {
    Icon(Icons.Default.Notifications, null)
}

DropdownMenu(
    expanded = showNotificaciones,
    onDismissRequest = { showNotificaciones = false },
    modifier = Modifier.width(300.dp)
) {
    notificaciones.forEach { promo ->
        DropdownMenuItem(
            text = { Column { Text(promo.producto); Text(promo.supermercado) } },
            onClick = { showNotificaciones = false; onPromocionesClick() }
        )
    }
}
''')

add_file_tag(doc, 'ui/screens/home/HomeScreen.kt  (ModalBottomSheet de perfil + AlertDialog de logout)')
add_code(doc, '''
var showSheet        by remember { mutableStateOf(false) }
var showLogoutDialog by remember { mutableStateOf(false) }
val sheetState = rememberModalBottomSheetState(skipPartiallyExpanded = true)

if (showSheet) {
    ModalBottomSheet(onDismissRequest = { showSheet = false }, sheetState = sheetState) {
        PerfilBottomSheetContent(
            onCerrarSesionClick = {
                scope.launch { sheetState.hide() }.invokeOnCompletion {
                    showSheet = false; showLogoutDialog = true
                }
            }
        )
    }
}

if (showLogoutDialog) {
    AlertDialog(
        onDismissRequest = { showLogoutDialog = false },
        title  = { Text(stringResource(R.string.perfil_cerrar_sesion)) },
        text   = { Text(stringResource(R.string.perfil_cerrar_sesion_confirmacion)) },
        confirmButton = {
            Button(onClick = { showLogoutDialog = false; onCerrarSesionClick() }) {
                Text(stringResource(R.string.eliminar_confirmar))
            }
        },
        dismissButton = {
            OutlinedButton(onClick = { showLogoutDialog = false }) { Text(stringResource(R.string.cancelar)) }
        }
    )
}
''')

add_file_tag(doc, 'ui/screens/compras/nueva/NuevaCompraScreen.kt  (DatePickerDialog + Dialog a medida)')
add_code(doc, '''
// Selector de fecha
if (mostrarCalendario) {
    DatePickerDialog(
        onDismissRequest = { mostrarCalendario = false },
        confirmButton = { /* ... */ }
    ) { /* DatePicker(...) */ }
}

// Diálogo a medida: elegir cámara o galería
if (showFotoDialog) {
    Dialog(onDismissRequest = { showFotoDialog = false }) {
        Surface(shape = RoundedCornerShape(20.dp)) {
            Column {
                // opción "Cámara"
                Box(Modifier.clickable {
                    showFotoDialog = false
                    permisoCamara.launch(Manifest.permission.CAMERA)
                })
                // opción "Galería"
                Box(Modifier.clickable {
                    showFotoDialog = false
                    galleryLauncher.launch("image/*")
                })
            }
        }
    }
}

// Diálogo a medida: agregar un supermercado nuevo a la lista
if (showAgregarDialog) {
    Dialog(onDismissRequest = { showAgregarDialog = false; nuevoSuperNombre = "" }) {
        Surface(shape = RoundedCornerShape(20.dp)) {
            Column { /* TextField con el nombre + botón "Agregar" */ }
        }
    }
}
''')

add_file_tag(doc, 'ui/screens/compras/historial/HistorialComprasScreen.kt  (FilterChip por mes)')
add_code(doc, '''
@Composable
private fun MesesRow(mesSel: Int?, onMesSelected: (Int?) -> Unit) {
    LazyRow {
        item {
            FilterChip(
                selected = mesSel == null,
                onClick  = { onMesSelected(null) },
                label    = { Text(stringResource(R.string.historial_filtro_todos)) }
            )
        }
        itemsIndexed(meses) { idx, nombre ->
            val mesNum = idx + 1
            FilterChip(
                selected = mesSel == mesNum,
                onClick  = { onMesSelected(if (mesSel == mesNum) null else mesNum) },
                label    = { Text(nombre) }
            )
        }
    }
}
''')

add_note(doc, 'Qué decir', '"Uso distintos tipos de menús y diálogos según el caso: DropdownMenu para una lista corta '
               'de notificaciones, ModalBottomSheet para el menú de perfil y para agregar un producto, AlertDialog '
               'para confirmar acciones que no se pueden deshacer (cerrar sesión, eliminar), Dialog a medida cuando '
               'necesito un diseño propio (elegir cámara o galería), DatePickerDialog para elegir la fecha de la '
               'compra, y FilterChip para los filtros de mes/categoría."')

doc.add_page_break()

# ================================================================== F. INTENTS
doc.add_heading('F. Intents — cámara/galería y compartir compra', level=1)

doc.add_heading('¿Qué es y para qué se usa?', level=2)
doc.add_paragraph(
    'Un Intent es la forma en que Android se comunica con otras apps o componentes del sistema. '
    'SuperAhorro usa Intents en dos lugares:'
)
bullet(doc, 'Para sacarle una foto al ticket (abre la app de Cámara) o elegir una imagen existente (abre la Galería), usando ActivityResultContracts y un FileProvider para que la cámara pueda escribir el archivo de forma segura.', bold_prefix='Nueva compra: ')
bullet(doc, 'El botón de compartir arma un Intent.ACTION_SEND con el resumen de la compra (texto plano) y abre el selector de apps del sistema (WhatsApp, Mail, etc.) con Intent.createChooser().', bold_prefix='Detalle de compra: ')

doc.add_heading('Demo en vivo', level=2)
numbered(doc, 'En "Nueva compra", tocar el botón de foto del ticket: se abre el diálogo "Cámara / Galería".')
numbered(doc, 'Elegir "Cámara": Android pide permiso de cámara (si es la primera vez) y luego abre la app de Cámara del sistema (Intent de captura).')
numbered(doc, 'Volver a la app: la foto sacada queda asociada al ticket.')
numbered(doc, 'Ir al detalle de una compra ya guardada y tocar el ícono de "Compartir" (arriba a la derecha): se abre el selector de apps de Android (WhatsApp, Gmail, etc.) con el resumen de la compra como texto.')

doc.add_heading('Código para mostrar', level=2)

add_file_tag(doc, 'ui/screens/compras/nueva/NuevaCompraScreen.kt  (cámara, galería y permiso)')
add_code(doc, '''
// Intent para abrir la cámara del sistema y guardar la foto en una Uri
val cameraLauncher = rememberLauncherForActivityResult(ActivityResultContracts.TakePicture()) { success ->
    if (success) cameraImageUri?.let { vm.setTicketUri(it) }
}

// Intent para elegir una imagen de la galería
val galleryLauncher = rememberLauncherForActivityResult(ActivityResultContracts.GetContent()) { uri ->
    uri?.let { vm.setTicketUri(it) }
}

// Pide el permiso de cámara antes de lanzar el Intent de la cámara
val permisoCamara = rememberLauncherForActivityResult(ActivityResultContracts.RequestPermission()) { granted ->
    if (granted) {
        val file = File(context.cacheDir, "ticket_${System.currentTimeMillis()}.jpg")
        val uri  = FileProvider.getUriForFile(context, "${context.packageName}.fileprovider", file)
        cameraImageUri = uri
        cameraLauncher.launch(uri)   // <- lanza el Intent de cámara
    }
}
''')

bullet(doc, 'TakePicture() / GetContent() son "contratos" de Android que ya saben armar el Intent correcto (cámara o selector de imágenes) y te devuelven el resultado en un callback.', bold_prefix='Línea clave 1: ')
bullet(doc, 'FileProvider.getUriForFile() crea una Uri "segura" (content://) para que la app de Cámara pueda escribir la foto sin darle acceso directo al sistema de archivos.', bold_prefix='Línea clave 2: ')
bullet(doc, 'RequestPermission() es el Intent del sistema para pedir el permiso CAMERA en tiempo de ejecución (Android 6+).', bold_prefix='Línea clave 3: ')

add_file_tag(doc, 'ui/screens/compras/detalle/DetalleCompraScreen.kt  (compartir compra: ACTION_SEND)')
doc.add_paragraph('Este es el "Compartir una compra por Intent" que pide la consigna como ejemplo concreto.')
add_code(doc, '''
IconButton(onClick = {
    compra?.let { c ->
        val texto = buildString {
            appendLine("Compra en ${c.supermercado}")
            appendLine("${c.fecha.format(fmt)} - ${c.hora}")
            c.productos.forEach { p ->
                appendLine("- ${p.nombre}  x${p.cantidad}  $${"%,.2f".format(p.subtotal)}")
            }
            appendLine("Total: $${"%,.2f".format(c.total)}")
            appendLine("Enviado desde Super Ahorro")
        }
        val intent = Intent(Intent.ACTION_SEND).apply {
            type = "text/plain"
            putExtra(Intent.EXTRA_TEXT, texto)
        }
        context.startActivity(Intent.createChooser(intent, "Compartir compra"))
    }
}) {
    Icon(Icons.Default.Share, null)
}
''')

add_note(doc, 'Qué decir', '"Tengo dos casos de Intents. Uno son los Intents de sistema para sacar fotos: pido el '
               'permiso de cámara con RequestPermission, genero una Uri segura con FileProvider y la paso al Intent '
               'de TakePicture; o uso GetContent para elegir una foto de la galería. El segundo caso es compartir '
               'una compra: armo un texto con el resumen de la compra y lo mando con Intent.ACTION_SEND dentro de '
               'un createChooser, que abre el selector de apps de Android para que el usuario elija con qué '
               'compartirlo, por ejemplo WhatsApp."')

doc.add_page_break()

# ================================================================== G. ETAPA FINAL
doc.add_heading('G. Etapa final — agregado diferencial', level=1)

doc.add_paragraph(
    'Esta sección agrupa las funcionalidades opcionales de la etapa final que ya están implementadas y '
    'funcionando con datos reales. Si el tiempo de la presentación es corto, alcanza con mostrar 1 o 2 '
    'de estas (recomendado: IA/OCR del ticket y biometría, porque son las más vistosas).'
)

doc.add_heading('G.1 — Carga automática del ticket con IA / OCR', level=2)
doc.add_paragraph(
    'Ya se mostró el código en la sección D (GroqRepository), pero acá se resalta el flujo completo: '
    'foto del ticket -> IA (Groq, modelo de visión) -> JSON con productos/fecha/hora/total -> autocompleta '
    'el formulario de Nueva compra.'
)
numbered(doc, 'En "Nueva compra", elegir/tomar la foto del ticket (Intents, sección F).')
numbered(doc, 'Tocar "Analizar con IA": se ve un indicador de carga.')
numbered(doc, 'La app completa sola el supermercado, fecha, hora, total y la lista de productos.')

add_file_tag(doc, 'ui/screens/compras/nueva/NuevaCompraViewModel.kt  (analizarTicketConIA)')
add_code(doc, '''
private val groqRepo = GroqRepository()

fun analizarTicketConIA(
    onResultado: (supermercado: String, fecha: String, hora: String, total: String, productos: List<Producto>) -> Unit
) {
    viewModelScope.launch {
        _analizando.value = true
        _errorIA.value    = null
        try {
            val bitmap: Bitmap = MediaStore.Images.Media.getBitmap(context.contentResolver, uri)
            when (val resultado = groqRepo.analizarTicket(bitmap)) {
                is GroqResult.Exito -> {
                    val t = resultado.ticket
                    _productos.value = t.productos
                    onResultado(t.supermercado, t.fecha, t.hora, t.total, t.productos)
                }
                is GroqResult.Error -> _errorIA.value = resultado.mensaje
            }
        } finally {
            _analizando.value = false
        }
    }
}
''')

add_note(doc, 'Qué decir', '"Esta es nuestra funcionalidad diferencial principal: en vez de tipear cada producto del '
               'ticket a mano, sacamos una foto y se la mandamos a un modelo de IA con visión (Groq), que devuelve '
               'un JSON estructurado con todos los productos, el supermercado, fecha, hora y total. El usuario '
               'puede revisar y corregir antes de guardar. Esto cubre tanto OCR como carga automática con IA."')

doc.add_heading('G.2 — Autenticación biométrica', level=2)
doc.add_paragraph(
    'En la pantalla de Login, si el dispositivo soporta huella digital o reconocimiento facial, aparece '
    'un botón adicional para loguearse con biometría en lugar de email/contraseña.'
)
numbered(doc, 'Iniciar sesión normalmente una vez (para que la sesión quede guardada en DataStore, sección A).')
numbered(doc, 'Cerrar sesión y volver a la pantalla de Login.')
numbered(doc, 'Tocar el botón de huella/biometría: el sistema pide la huella o el rostro.')
numbered(doc, 'Al validar, la app entra directo al Home con el último usuario logueado.')

add_file_tag(doc, 'ui/screens/login/LoginScreen.kt')
add_code(doc, '''
val puedeUsarBiometria = remember {
    BiometricManager.from(context).canAuthenticate(
        BiometricManager.Authenticators.BIOMETRIC_STRONG or
        BiometricManager.Authenticators.BIOMETRIC_WEAK
    ) == BiometricManager.BIOMETRIC_SUCCESS
}

fun lanzarBiometria() {
    val activity = context as FragmentActivity
    val prompt = BiometricPrompt(
        activity,
        ContextCompat.getMainExecutor(context),
        object : BiometricPrompt.AuthenticationCallback() {
            override fun onAuthenticationSucceeded(result: BiometricPrompt.AuthenticationResult) {
                vm.loginConBiometria(onLoginExitoso)
            }
        }
    )
    val promptInfo = BiometricPrompt.PromptInfo.Builder()
        .setTitle(context.getString(R.string.login_biometria_titulo))
        .setNegativeButtonText(context.getString(R.string.cancelar))
        .setAllowedAuthenticators(
            BiometricManager.Authenticators.BIOMETRIC_STRONG or
            BiometricManager.Authenticators.BIOMETRIC_WEAK
        )
        .build()
    prompt.authenticate(promptInfo)
}

// El botón solo se muestra si el dispositivo soporta biometría
if (puedeUsarBiometria) {
    OutlinedButton(onClick = { lanzarBiometria() }) { /* ícono de huella */ }
}
''')

add_file_tag(doc, 'data/repository/AuthRepository.kt  (loginConBiometria)')
add_code(doc, '''
suspend fun loginConBiometria(): AuthResult {
    val userId = session.userId.first()                 // id guardado en DataStore (sección A)
    val usuario = dao.getById(userId)                    // se busca en Room (sección B)
    return if (usuario != null) AuthResult.Exito(usuario)
    else AuthResult.Error("No hay una sesión guardada")
}
''')

add_note(doc, 'Qué decir', '"La biometría no reemplaza el login: valida la huella o el rostro con BiometricPrompt y, '
               'si es correcto, reutiliza el id de usuario que ya tenemos guardado en DataStore para recuperar sus '
               'datos de Room. El botón solo aparece si BiometricManager detecta que el dispositivo tiene huella o '
               'reconocimiento facial configurado."')

doc.add_heading('G.3 — Notificaciones (in-app)', level=2)
doc.add_paragraph(
    'En el Home, el ícono de campana muestra un punto rojo cuando hay ofertas nuevas, y al tocarlo se '
    'abre el DropdownMenu (sección E) con las 5 mejores promociones traídas en vivo desde la API '
    '(sección D), ordenadas por mayor descuento.'
)
numbered(doc, 'Abrir el Home y mostrar el punto rojo sobre la campana (indicador de "nueva oferta").')
numbered(doc, 'Tocar la campana: se despliega la lista de ofertas con producto, supermercado y descuento.')
numbered(doc, 'Tocar "Ver todas": navega a la pantalla de Promociones.')

add_file_tag(doc, 'ui/screens/home/HomeViewModel.kt')
add_code(doc, '''
private val promoRepo = PromocionesRepository()
private val _notificaciones = MutableStateFlow<List<Promocion>>(emptyList())
val notificaciones = _notificaciones.asStateFlow()

private fun cargarNotificaciones() {
    viewModelScope.launch {
        val resultado = promoRepo.buscarPromocionesArgentina()
        if (resultado is PromocionesResult.Exito) {
            // Ordena por mayor descuento primero y se queda con las 5 mejores
            _notificaciones.value = resultado.promociones
                .sortedBy { promo ->
                    val original = promo.precioSinDescuento
                    if (original != null && original > 0) (promo.precio / original) else 1.0
                }
                .take(5)
        }
    }
}
''')

add_note(doc, 'Qué decir', '"Las notificaciones no son inventadas: son las mismas promociones reales que trae '
               'PromocionesRepository, ordenadas por mayor descuento, y se muestran como avisos de \\"nueva oferta\\" '
               'en el Home. Es una forma de avisar al usuario apenas entra a la app, sin que tenga que ir manualmente '
               'a la pantalla de Promociones."')

doc.add_heading('G.4 — Filtros avanzados', level=2)
doc.add_paragraph(
    'La pantalla de Historial permite filtrar las compras por año y por mes; la pantalla de Listado '
    'permite buscar por texto (supermercado o nombre de producto).'
)
numbered(doc, 'En "Historial de compras", cambiar el año con las flechas del selector.')
numbered(doc, 'Tocar un mes (chip) para ver solo las compras de ese mes; tocar "Todos" para volver a ver todo el año.')
numbered(doc, 'En "Mis compras" (Listado), escribir en el buscador el nombre de un supermercado o producto y mostrar cómo se filtra la lista en tiempo real.')

add_file_tag(doc, 'ui/screens/compras/historial/HistorialComprasViewModel.kt')
add_code(doc, '''
private val _mesSel  = MutableStateFlow<Int?>(null)
private val _anioSel = MutableStateFlow(LocalDate.now().year)

fun seleccionarMes(mes: Int?) {
    _mesSel.value = mes
    aplicarFiltros()
}

fun seleccionarAnio(anio: Int) {
    _anioSel.value = anio
    _mesSel.value  = null
    aplicarFiltros()
}

private fun aplicarFiltros() {
    val mes  = _mesSel.value
    val anio = _anioSel.value
    val resultado = _todasLasCompras.value.filter { c ->
        c.fecha.year == anio && (mes == null || c.fecha.monthValue == mes)
    }
    // ... se agrupa por mes y se calcula el total de cada grupo ...
}
''')

add_note(doc, 'Qué decir', '"En Historial el usuario puede navegar por año y filtrar por mes con los chips; el '
               'ViewModel mantiene el mes y año seleccionados como estado y vuelve a calcular la lista filtrada '
               'y los totales cada vez que cambian. En el listado de compras también hay un buscador que filtra '
               'por supermercado o producto en tiempo real."')

doc.add_page_break()

# ================================================================== NOTAS HONESTAS
doc.add_heading('Notas honestas: qué falta o está parcial', level=1)

doc.add_paragraph(
    'Si el profesor pregunta por alguno de estos puntos, mejor responder con honestidad que el resto del '
    'trabajo ya cubre sobradamente los requisitos obligatorios, y que estos quedan como próximos pasos.'
)

add_warning(doc, 'Exportación de datos (CSV/PDF)',
            'En Configuración existe la fila "Exportar" (CSV / PDF), pero todavía no tiene una acción '
            'asociada (es solo informativa/visual). Queda como mejora a futuro.')

add_warning(doc, 'Comparativa de precios entre supermercados',
            'Está cubierta de forma parcial: Promociones muestra ofertas de Carrefour y Changomás juntas, '
            'y Estadísticas muestra el porcentaje de gasto por supermercado. No hay todavía una comparación '
            'directa "mismo producto en distintos supermercados".')

add_warning(doc, 'Chat para consultas sobre historial de compras',
            'No está implementado. Podría construirse a futuro reutilizando la integración con la API de '
            'Groq que ya existe para el análisis de tickets.')

add_note(doc, 'Qué decir si preguntan', '"De los 8 opcionales de etapa final, ya tenemos 4 completos con datos '
               'reales (IA/OCR, biometría, notificaciones y filtros), 1 parcial (comparativa entre supermercados) '
               'y 3 quedan como próximos pasos (exportación, chat y profundizar la comparativa). Priorizamos que '
               'todo lo que mostramos funcione de punta a punta con datos reales, en vez de tener muchas '
               'funcionalidades a medias."')

doc.add_page_break()

# ================================================================== TABLA RESUMEN
doc.add_heading('Tabla resumen (cheat sheet)', level=1)

make_table(
    doc,
    ['Requisito', 'Archivos clave', 'Idea en una frase'],
    [
        ('A. Persistencia local (DataStore)',
         'SessionManager.kt\nThemeviewmodel.kt\nAuthRepository.kt (login)',
         'Persiste sesión y preferencia de tema con Flow + corrutinas, sin bloquear la UI.'),
        ('B. Base de datos local (Room)',
         'AppDatabase.kt\nUsuarioEntity / CompraEntity / ProductoEntity\nCompraDao / CompraRepository.kt',
         'Tablas relacionadas (usuarios, compras, productos) con borrado en cascada y lectura reactiva vía Flow.'),
        ('C. Corrutinas',
         'viewModelScope.launch, withContext(IO),\nFlow.collect, .first(), mapLatest',
         'Toda lectura/escritura (BD, DataStore, red) corre en corrutinas, sin bloquear la interfaz.'),
        ('D. Networking',
         'PromocionesRepository.kt\nPromocionesViewModel.kt\nGroqRepository.kt',
         'OkHttp + JSON contra APIs públicas reales (VTEX de supermercados e IA de Groq).'),
        ('E. Menús y diálogos',
         'HomeScreen.kt (Dropdown, BottomSheet, AlertDialog)\nNuevaCompraScreen.kt (Dialog, DatePicker)\nHistorialComprasScreen.kt (FilterChip)',
         'DropdownMenu, ModalBottomSheet, AlertDialog, Dialog a medida, DatePickerDialog y FilterChip, cada uno para su caso.'),
        ('F. Intents',
         'NuevaCompraScreen.kt (cámara/galería)\nDetalleCompraScreen.kt (compartir)',
         'TakePicture/GetContent + FileProvider para fotos, y ACTION_SEND + createChooser para compartir una compra.'),
        ('G. Etapa final',
         'GroqRepository.kt (IA/OCR)\nLoginScreen.kt (biometría)\nHomeViewModel.kt (notificaciones)\nHistorialComprasViewModel.kt (filtros)',
         'IA/OCR de tickets, login biométrico, notificaciones de ofertas reales y filtros por año/mes.'),
    ],
    font_size=9
)

doc.add_page_break()

# ================================================================== PREGUNTAS
doc.add_heading('Posibles preguntas y cómo responderlas', level=1)

qas = [
    ('¿Por qué DataStore y no SharedPreferences?',
     'DataStore es la recomendación actual de Google: es asincrónico (no bloquea el hilo principal), '
     'usa corrutinas/Flow y es type-safe con las Preferences keys.'),
    ('¿Qué pasa si no hay internet al pedir las promociones o analizar un ticket?',
     'Cada supermercado se consulta dentro de un try/catch independiente; si uno falla, se sigue con el '
     'resto. Si fallan todos, el repositorio devuelve PromocionesResult.Error / GroqResult.Error y la '
     'pantalla muestra un mensaje en vez de romperse.'),
    ('¿Por qué Room y no manejar SQLite directo o Firebase?',
     'Room evita escribir SQL a mano para lo básico, valida las consultas en tiempo de compilación y expone '
     'Flow para que la UI se actualice sola cuando cambian los datos.'),
    ('¿Qué relación hay entre compras y productos?',
     'Una compra (CompraEntity) tiene muchos productos (ProductoEntity), relacionados por compraId con una '
     'ForeignKey con onDelete = CASCADE: si se borra la compra, se borran sus productos.'),
    ('¿Cómo se actualiza la UI cuando cambian los datos en Room?',
     'El DAO devuelve un Flow<List<...>>. El ViewModel hace collect de ese Flow y actualiza el UiState; '
     'Compose observa ese UiState con collectAsStateWithLifecycle, así que la pantalla se redibuja sola.'),
    ('¿Dónde se guarda la API key de Groq y por qué no está en el código?',
     'Se lee desde BuildConfig.GROQ_API_KEY, que se genera a partir de local.properties (un archivo que no '
     'se sube al repositorio), para no exponer la clave.'),
    ('¿Cómo funciona el permiso de cámara?',
     'Se usa ActivityResultContracts.RequestPermission() para pedir Manifest.permission.CAMERA en tiempo de '
     'ejecución. Solo si el usuario lo concede, se genera la Uri con FileProvider y se lanza el Intent de '
     'cámara (TakePicture).'),
    ('¿Por qué usar FileProvider en vez de pasar la ruta del archivo directo?',
     'Desde Android 7 (API 24) no se puede compartir una ruta "file://" directa entre apps por seguridad. '
     'FileProvider genera una Uri "content://" temporal y segura para que la app de Cámara pueda escribir '
     'la foto.'),
    ('¿Qué pasa si el celular no tiene huella/biometría configurada?',
     'BiometricManager.canAuthenticate() devuelve algo distinto de BIOMETRIC_SUCCESS, y el botón de '
     'biometría simplemente no se muestra (puedeUsarBiometria = false). El login normal con email/contraseña '
     'sigue funcionando igual.'),
    ('¿De dónde salen las "notificaciones"?',
     'Son las mismas promociones reales que devuelve PromocionesRepository, ordenadas por mayor descuento y '
     'recortadas a las 5 mejores. No son notificaciones push del sistema, son un aviso dentro de la app '
     '(menú desplegable en el Home).'),
    ('¿Qué pasa si cambia el modelo de la base de datos?',
     'Se sube el número de "version" en @Database. En este proyecto se usa fallbackToDestructiveMigration(), '
     'pensado para la etapa de desarrollo: recrea la base si la versión cambia.'),
]

for q, a in qas:
    p = doc.add_paragraph()
    r = p.add_run('P: ')
    r.bold = True
    r.font.color.rgb = RGBColor.from_string('C0392B')
    p.add_run(q).bold = True
    p2 = doc.add_paragraph()
    r2 = p2.add_run('R: ')
    r2.bold = True
    r2.font.color.rgb = RGBColor.from_string('2E75B6')
    p2.add_run(a)
    p2.paragraph_format.space_after = Pt(10)

doc.add_page_break()

# ================================================================== TIPS FINALES
doc.add_heading('Tips finales para el día de la presentación', level=1)

bullet(doc, 'Empezá siempre por la app funcionando, después mostrá el código: es más fácil de seguir para quien evalúa.')
bullet(doc, 'No leas el código línea por línea: señalá 2-3 puntos clave por archivo (los marcados en esta guía) y explicalos con tus palabras.')
bullet(doc, 'Si algo falla en vivo (por ejemplo, la API de promociones por la red del aula), quedate tranquilo: mostrá el código, explicá qué hace, y mencioná que funciona con conexión (podés tener una captura de pantalla de respaldo).')
bullet(doc, 'Remarcá las conexiones entre secciones: el login usa Room (validar usuario) y DataStore (guardar sesión); el historial usa Room y el id de usuario que viene de DataStore; Nueva compra combina Room, Intents (foto) y Networking (IA); las notificaciones del Home reusan la misma API que Promociones.')
bullet(doc, 'Si te preguntan algo que no sabés con certeza, está bien decir "lo tengo que revisar" antes que inventar una respuesta. Para los puntos pendientes (exportación, chat), está bien decir que quedaron como próximos pasos.')

doc.add_paragraph()
final = doc.add_paragraph()
final.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = final.add_run('¡Éxitos en la presentación del viernes!')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor.from_string('1F4E79')

# ================================================================== GUARDAR
out_path = r'C:\Users\Francisco Bossio\AndroidStudioProjects\SuperAhorro\Guia_Presentacion_SuperAhorro.docx'
doc.save(out_path)
print('OK ->', out_path)
